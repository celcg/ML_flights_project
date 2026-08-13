"""Build the one-page, two-column executive summary for the flight-delay project."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "doc" / "Summary_Report_aportaciones_modelo_vuelos.docx"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
PALE_BLUE = "DCE6F1"
PALE_GREY = "F2F4F7"
MID_GREY = "667085"
WHITE = "FFFFFF"
GREEN = "1B7F5A"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D0D5DD", size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_columns(section, number=2, space_twips=420) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(number))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    p_pr.append(node)


def set_keep_together(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepLines")
    p_pr.append(node)


def add_heading(doc, text: str) -> None:
    p = doc.add_paragraph(style="Summary Heading")
    p.add_run(text)
    keep_with_next(p)


def add_bullet(doc, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="Summary Bullet")
    bullet = p.add_run("• ")
    bullet.bold = True
    bullet.font.color.rgb = RGBColor.from_string(BLUE)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    set_keep_together(p)


def add_numbered(doc, number: int, text: str) -> None:
    p = doc.add_paragraph(style="Summary Body")
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r = p.add_run(f"{number}. ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(text)
    set_keep_together(p)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string("243142")
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    heading = doc.styles.add_style("Summary Heading", WD_STYLE_TYPE.PARAGRAPH)
    heading.font.name = "Calibri"
    heading.font.size = Pt(12)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(BLUE)
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(4)

    body = doc.styles.add_style("Summary Body", WD_STYLE_TYPE.PARAGRAPH)
    body.font.name = "Calibri"
    body.font.size = Pt(9.4)
    body.font.color.rgb = RGBColor.from_string("243142")
    body.paragraph_format.space_after = Pt(3)
    body.paragraph_format.line_spacing = 1.04

    bullet = doc.styles.add_style("Summary Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(9.3)
    bullet.font.color.rgb = RGBColor.from_string("243142")
    bullet.paragraph_format.left_indent = Cm(0.42)
    bullet.paragraph_format.first_line_indent = Cm(-0.28)
    bullet.paragraph_format.space_after = Pt(2.8)
    bullet.paragraph_format.line_spacing = 1.02
    bullet._element.get_or_add_pPr().append(OxmlElement("w:contextualSpacing"))


def add_metric_strip(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(6.15)
    table.columns[1].width = Cm(6.15)
    table.columns[2].width = Cm(6.15)
    cells = table.rows[0].cells
    items = (
        ("RMSE GLOBAL · 25%", "15,49 → 14,75 min", "−4,8% frente al baseline"),
        ("MAE VUELOS >15 MIN", "21,66 → 18,20 min", "−16,0% frente al baseline"),
        ("CRITERIO COMBINADO", "14,05 min", "Ridge: mejor equilibrio"),
    )
    for idx, (label, value, note) in enumerate(items):
        cell = cells[idx]
        shade(cell, PALE_BLUE if idx != 2 else "DDF2EA")
        set_cell_margins(cell, top=80, start=110, bottom=75, end=110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7.8)
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE if idx != 2 else GREEN)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        r2.bold = True
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(note)
        r3.font.size = Pt(7.4)
        r3.font.color.rgb = RGBColor.from_string(MID_GREY)
    set_table_borders(table, color=WHITE, size="8")


def add_results_table(doc: Document) -> None:
    rows = (
        ("Global MAE", "10,13", "9,89", "−2,3%"),
        ("Global RMSE", "15,49", "14,75", "−4,8%"),
        ("MAE >15 min", "21,66", "18,20", "−16,0%"),
        ("RMSE >15 min", "29,11", "25,86", "−11,2%"),
    )
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(2.7), Cm(1.35), Cm(1.25), Cm(1.25))
    for col, width in zip(table.columns, widths):
        col.width = width
    headers = ("Métrica (min)", "Baseline", "Ridge", "Mejora")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, DARK_BLUE)
        set_cell_margins(cell, top=55, start=55, bottom=55, end=55)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(7.2)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            shade(cell, PALE_GREY if row_idx % 2 else WHITE)
            set_cell_margins(cell, top=52, start=55, bottom=52, end=55)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            r.font.size = Pt(7.8)
            if col_idx == 3:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(GREEN)
    set_table_borders(table)


def add_callout(doc: Document, title: str, text: str, fill=PALE_BLUE) -> None:
    # A shaded paragraph, rather than a table, respects native column width in
    # LibreOffice and Word. Full-width table defaults can otherwise overflow.
    p = doc.add_paragraph(style="Summary Body")
    p.paragraph_format.left_indent = Cm(0.08)
    p.paragraph_format.right_indent = Cm(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.04
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), BLUE if fill == PALE_BLUE else "E4A11B")
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run("\n")
    p.add_run(text)
    set_keep_together(p)


def add_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Validación temporal: diciembre de 2022 · 29.315 vuelos · test de marzo de 2023 intacto"
    )
    r.font.name = "Calibri"
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor.from_string(MID_GREY)


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    # Named layout override: the standard business brief is compacted to fit the
    # user-requested single A4 sheet while retaining its typography and blue palette.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Cm(0.65)
    section.bottom_margin = Cm(0.65)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.header_distance = Cm(0.25)
    section.footer_distance = Cm(0.28)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(1)
    r = title.add_run("SUMMARY REPORT")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(5)
    r = subtitle.add_run("Predicción del retraso de llegada una hora antes de la salida")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r2 = subtitle.add_run("  |  Estado a 13 de agosto de 2026")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor.from_string(MID_GREY)

    add_metric_strip(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = section.top_margin
    body_section.bottom_margin = section.bottom_margin
    body_section.left_margin = section.left_margin
    body_section.right_margin = section.right_margin
    body_section.header_distance = section.header_distance
    body_section.footer_distance = section.footer_distance
    body_section.footer.is_linked_to_previous = True
    set_columns(body_section, number=2, space_twips=420)

    add_heading(doc, "Mayores aportaciones")
    add_bullet(
        doc,
        "Horizonte operativo T−60: solo usa información disponible una hora antes; la auditoría detectó cero eventos futuros.",
    )
    add_bullet(
        doc,
        "Variables operativas: actividad y retraso observado por origen, destino, ruta y operador en ventanas de 1 y 24 horas; se descartó 6 h tras la ablación.",
    )
    add_bullet(
        doc,
        "Rotación de aeronave: retraso del vuelo anterior, tiempo desde su llegada y disponibilidad del avión.",
    )
    add_bullet(
        doc,
        "Calendario y vuelo: duración programada, hora/día cíclicos, mes y nivel de vuelo solicitado.",
    )
    add_bullet(
        doc,
        "Categorías: one-hot solo para baja cardinalidad; AC Type raro → OTHER y hashing para aeronave, aeropuertos y operador.",
    )

    add_heading(doc, "Modelos probados")
    add_bullet(doc, "Baselines: mediana global, ruta y ruta+aerolínea con fallbacks.")
    add_bullet(doc, "Ridge con variables sin transformar, log y Yeo–Johnson.")
    add_bullet(doc, "Random Forest, Gradient-Boosted Trees, XGBoost y CatBoost.")
    add_bullet(doc, "Ensemble Ridge + CatBoost con variables operativas T−60.")

    add_heading(doc, "Ridge frente al baseline histórico")
    add_results_table(doc)
    p = doc.add_paragraph(style="Summary Body")
    p.paragraph_format.space_before = Pt(2)
    p.add_run("Comparación homogénea: ").bold = True
    p.add_run("ambos ajustados con el mismo 25% determinista de train y evaluados en las mismas 29.315 filas.")

    # Force the following content into the second native page column.
    breaker = doc.add_paragraph()
    breaker.paragraph_format.space_after = Pt(0)
    breaker.add_run().add_break(WD_BREAK.COLUMN)

    add_heading(doc, "¿Por qué Ridge es la mejor elección actual?")
    add_bullet(
        doc,
        "Selección alineada al 10%: MAE global y MAE de vuelos con más de 15 min pesan al 50%; Ridge obtiene el mejor criterio combinado, 14,07 min.",
    )
    add_bullet(
        doc,
        "En esa comparación al 10%, CatBoost gana ligeramente en promedio global (MAE 9,54; RMSE 14,67), pero Ridge predice mejor los retrasos elevados (MAE 18,21 frente a 19,63).",
    )
    add_bullet(
        doc,
        "Es estable, regularizado y eficiente con matrices dispersas y hashing: una ventaja decisiva con la RAM disponible.",
    )
    add_bullet(
        doc,
        "Su lectura es clara: las señales operativas recientes y la rotación aportan más que aumentar por sí sola la complejidad del algoritmo.",
    )

    add_callout(
        doc,
        "Interpretación ejecutiva",
        "Con 25% de train, Ridge reduce el RMSE global en 0,74 minutos (4,8%) y el MAE de retrasos elevados en 3,46 minutos (16,0%). El mayor valor está en detectar mejor los casos operativamente difíciles.",
    )

    add_heading(doc, "Siguientes pasos")
    add_numbered(
        doc,
        1,
        "25% completado: 613.884 vuelos de train. El criterio combinado mejora solo 0,027 min frente al 10%; no continuar aún a 50%/100%.",
    )
    add_numbered(
        doc,
        2,
        "Priorizar más meses/años y nuevas señales operativas; el límite actual no parece ser la cantidad de filas dentro de los mismos meses.",
    )
    add_numbered(
        doc,
        3,
        "Validar estabilidad por mes, aeropuerto, ruta y nivel de retraso; mantener marzo de 2023 como prueba final intacta.",
    )
    add_numbered(
        doc,
        4,
        "Ampliar meses/años antes de añadir meteorología; aporta mejor cobertura de estacionalidad y eventos poco frecuentes.",
    )
    add_numbered(
        doc,
        5,
        "Después, evaluar meteorología, modelo en dos etapas (retrasado/no retrasado + minutos) y LightGBM/SynapseML como trabajo futuro.",
    )

    add_callout(
        doc,
        "Decisión de escalado: no continuar todavía",
        "Ridge 25% mejora el MAE y RMSE globales solo 0,049 min frente al 10%. Es una señal positiva, pero el beneficio práctico no compensa aún el coste de entrenar al 50%/100%. Test no leído.",
        fill="DDF2EA",
    )

    doc.core_properties.title = "Summary Report — Predicción de retraso de llegada T−60"
    doc.core_properties.subject = "Aportaciones, resultados y siguientes pasos"
    doc.core_properties.author = "Proyecto ML Flights"
    doc.core_properties.keywords = "flight delay, Ridge, CatBoost, T-60, MAE, RMSE"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
