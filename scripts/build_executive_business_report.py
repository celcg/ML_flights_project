"""Build the curated executive aviation business report."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from scripts.build_full_business_report import (
    BLUE,
    DARK_BLUE,
    EDA,
    EXTRA,
    GREEN,
    LIGHT_BLUE,
    ORANGE,
    ROOT,
    add_body,
    add_callout,
    add_data_table,
    add_figure,
    add_title_page,
    configure_document,
    f,
    page_heading,
    rows,
    section_break,
)


OUTPUT = ROOT / "doc" / "European_Scheduled_Aviation_Business_Report_WITH_CASE_STUDIES_2026-08-14.docx"


def _rank(data, key: str, descending: bool = False, n: int = 5):
    return sorted(data, key=lambda row: f(row, key), reverse=descending)[:n]


def build() -> Path:
    tables = EXTRA / "tables"
    countries = rows(tables / "country_statistics_min_1000_movements.csv")
    origin_best = rows(tables / "most_reliable_origin_airports.csv")
    origin_worst = rows(tables / "most_problematic_origin_airports.csv")
    destination_best = rows(tables / "most_reliable_destination_airports.csv")
    destination_worst = rows(tables / "most_problematic_destination_airports.csv")
    carrier_best = rows(tables / "most_reliable_operating_carriers.csv")
    carrier_worst = rows(tables / "most_problematic_operating_carriers.csv")
    carrier_duration_best = rows(tables / "best_operating_carriers_by_relative_delay.csv")
    carrier_duration_worst = rows(tables / "worst_operating_carriers_by_relative_delay.csv")
    ryanair_wizz = rows(tables / "ryanair_wizz_shared_route_summary.csv")
    madrid_periods = rows(tables / "madrid_milan_period_performance.csv")
    madrid_operators = rows(tables / "madrid_milan_operator_performance.csv")
    madrid_hours = rows(tables / "madrid_milan_hour_performance.csv")
    madrid_extremes = rows(tables / "madrid_milan_largest_observed_arrival_delays.csv")
    pressure = rows(tables / "within_airport_pressure_band_performance.csv")
    weekdays = rows(tables / "weekday_network_performance.csv")
    day_hours = rows(tables / "day_hour_otp15_performance.csv")
    problematic_routes = rows(EDA / "tables" / "least_reliable_routes.csv")

    countries_best = _rank(countries, "otp15_ci_low_pct", descending=True)
    countries_worst = _rank(countries, "otp15_ci_high_pct")

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Executive summary", level=1)
    add_callout(
        doc,
        "Network result",
        "82.2% of 5.38 million eligible operated flights arrived within 15 minutes of schedule. OTP15 is the sole decision metric in this version.",
        GREEN,
        DARK_BLUE,
    )
    add_body(
        doc,
        "This report is an operational screening tool. It identifies large, repeated differences by airport, country or territory, operating carrier, traffic pressure, weekday and hour. It does not assign responsibility because route mix, schedule design and external disruption affect raw comparisons.",
    )
    add_body(
        doc,
        "The clearest network signal is traffic pressure. The share arriving more than 15 minutes late rises from 12.0% in each airport's lowest traffic quartile to 21.1% in its peak quartile. This 9.1-point gap exceeds the agreed three-percentage-point business threshold, although it remains an association rather than causal proof.",
    )
    add_body(
        doc,
        "Weekday differences are statistically detectable because the sample is very large, but no weekday differs from the rest of the network by the agreed three-point threshold. Tuesday is descriptively strongest at 84.3% OTP15 and Saturday weakest at 80.3%; the practical focus should be day-by-hour windows, not weekday alone.",
    )
    add_callout(
        doc,
        "Management focus",
        "Use high-volume maps and Wilson rankings to select investigations; monitor the ten busiest origin airports by period; and control route, airport, duration and hour before publishing a fair airline league table.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    section_break(doc, "1. Methodology and decision rules")
    doc.add_heading("Scope and denominator", level=2)
    add_body(
        doc,
        "The analysis uses nine observed monthly snapshots: June, September and December 2021; March, June, September and December 2022; and March and June 2023. They are not consecutive monthly data. Only regular scheduled flights (ICAO Flight Type S) are included, and arrival OTP15 requires an observed operated arrival. The source has no complete passenger or cancellation denominator.",
    )
    add_body(
        doc,
        "Quality rules exclude impossible early delays below -120 minutes, flight levels outside 0-500 and non-positive flown distance when those measurements are present. Null physical measurements are retained when the KPI does not depend on them.",
    )
    add_body(
        doc,
        "Executive eligibility is stricter than exploratory filtering. Airports require at least 200 flights and presence in three snapshots; operating carriers require 1,000 flights; routes require 20 flights and three snapshots; countries or territories require 1,000 historical endpoint movements. A jurisdiction represented by one airport remains visible but is marked as limited representation.",
    )
    doc.add_heading("Why Wilson, p-tests and a practical threshold", level=2)
    add_body(
        doc,
        "Wilson 95% intervals show uncertainty around OTP15 and behave well near 0% and 100%. Rankings use conservative Wilson bounds, which prevents a small group with a lucky point estimate from automatically leading the table.",
    )
    add_body(
        doc,
        "Each eligible group is compared with the rest of the network using a two-proportion z-test. Benjamini-Hochberg adjustment controls the false-discovery rate across many comparisons. Because millions of flights make tiny differences statistically significant, a result is decision-relevant only when the adjusted p-value is below 0.05 and the absolute OTP15 difference is at least three percentage points.",
    )
    add_body(
        doc,
        "Trend lines use the nine snapshots, require at least six observed periods and estimate a June 2021-to-June 2023 change. The top-10 airport alert matrix compares each period with that airport's other periods, avoiding a benchmark that contains the value being assessed. These are monitoring signals, not causal estimates.",
    )
    add_callout(
        doc,
        "Delay definition used in the report",
        "Unless a figure explicitly says departure delay, OTP15 and relative-delay burden use arrival delay: ACTUAL ARRIVAL TIME minus FILED ARRIVAL TIME. OTP15 means that this arrival difference is at most 15 minutes. Statistical significance tests whether a difference is likely to be noise; the three-point threshold tests whether it is large enough to matter operationally. Both conditions are required.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    add_figure(
        doc,
        EDA / "figures" / "statistical_method_explainer.png",
        "Figure 1. How rates, uncertainty, hypothesis tests and practical importance work together.",
        "Decision rule: adjusted p < 0.05 and absolute OTP15 difference at least 3 percentage points.",
        15.5,
    )
    section_break(doc, "2. Airports")
    add_figure(
        doc,
        EDA / "figures" / "origin_airport_reliability_rankings.png",
        "Figure 2. Eligible origin-airport OTP15 rankings with 95% Wilson intervals.",
        "Eligibility: at least 200 flights and three observed periods; ranking uses Wilson bounds.",
        15.8,
    )
    add_figure(
        doc,
        EDA / "figures" / "destination_airport_reliability_rankings.png",
        "Figure 3. Eligible destination-airport OTP15 rankings with 95% Wilson intervals.",
        "Results describe observed outcomes and are not adjusted causal airport effects.",
        15.8,
    )
    add_body(
        doc,
        f"The three most problematic eligible origins are {origin_worst[0]['airport']} ({f(origin_worst[0], 'arrival_otp15_pct'):.1f}% OTP15), {origin_worst[1]['airport']} ({f(origin_worst[1], 'arrival_otp15_pct'):.1f}%) and {origin_worst[2]['airport']} ({f(origin_worst[2], 'arrival_otp15_pct'):.1f}%). These now match the plotted top three. Their long-haul geography requires a route-direction and schedule-time audit before attributing the result to local airport operations.",
    )
    add_body(
        doc,
        f"The three most reliable eligible origins are {origin_best[0]['airport']}, {origin_best[1]['airport']} and {origin_best[2]['airport']}. The leading destinations are {destination_best[0]['airport']}, {destination_best[1]['airport']} and {destination_best[2]['airport']}; the lowest-ranked eligible destination is {destination_worst[0]['airport']} at {f(destination_worst[0], 'arrival_otp15_pct'):.1f}% OTP15.",
    )

    page_heading(doc, "Airport geography and monitoring", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "top200_airport_otp15_geographic_maps.png",
        "Figure 4. Top-200 airport volume and OTP15 on a grayscale geographic basemap.",
        "The upper panel focuses on Europe and clips the far east of Russia and Turkey; the lower panel checks global coverage. Bubble area is movement volume.",
        17.2,
    )
    page_heading(doc, "Airport delay burden beyond OTP15", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "top200_airport_relative_delay_comparison.png",
        "Figure 4B. The same high-volume destination airports compared by OTP15 and relative delay burden.",
        "Relative burden is total positive arrival-delay minutes divided by total scheduled flight minutes. Bubble area represents observed arrivals; the original OTP15 geographic map is retained on the preceding page.",
        17.2,
    )
    relative_map = rows(tables / "top200_destination_relative_delay.csv")
    visible_europe = [
        row
        for row in relative_map
        if -25 <= f(row, "longitude") <= 40 and 30 <= f(row, "latitude") <= 72
    ]
    highest_burden = _rank(
        visible_europe,
        "relative_delay_burden_pct",
        descending=True,
        n=3,
    )
    add_body(
        doc,
        "The highest visible relative-delay burdens are "
        + ", ".join(
            f"{row['airport']} ({f(row, 'relative_delay_burden_pct'):.1f}%)"
            for row in highest_burden
        )
        + ". This metric is not the percentage of flights delayed: it measures positive delay minutes relative to scheduled flying minutes.",
    )
    add_callout(
        doc,
        "Interpret with OTP15",
        "OTP15 measures frequency; relative delay burden adds severity scaled to flight duration. A high value in both panels is a stronger investigation signal, but long-haul schedule times, timezones and date rollover must be audited before attributing performance.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    doc.add_page_break()
    add_figure(
        doc,
        EXTRA / "figures" / "top10_origin_airport_period_alerts.png",
        "Figure 5. Period alerts for the ten highest-volume origin airports.",
        "Each snapshot is compared with the airport's other snapshots; an asterisk requires adjusted significance and at least a 3-point difference.",
        17.0,
    )
    add_callout(
        doc,
        "How to use the alert matrix",
        "Red means material deterioration, blue material improvement, green a difference inside the three-point tolerance and grey an inconclusive result. Limiting the matrix to ten airports keeps it operationally readable.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    section_break(doc, "3. Countries and territories")
    add_body(
        doc,
        "Country-or-territory eligibility is at least 1,000 historical endpoint movements. A flight can contribute one origin and one destination endpoint. ISO-coded territories represented by one airport are retained, as requested, but marked as limited representation.",
    )
    add_figure(
        doc,
        EXTRA / "figures" / "country_reliability_min_1000_movements.png",
        "Figure 6. Best and worst eligible country-or-territory OTP15 results.",
        "Ranking uses Wilson bounds; limited marks identify a one-airport representation.",
    )
    add_figure(
        doc,
        EXTRA / "figures" / "country_otp15_geographic_map.png",
        "Figure 7. Eligible country-or-territory OTP15 and represented movement volume in Europe.",
        "Bubble position is the mean coordinate of represented airports; the basemap is geographic context only.",
    )
    add_data_table(
        doc,
        ["Country or territory", "Airports", "Movements", "OTP15", "Representation"],
        [
            [
                r["country"],
                int(f(r, "airports")),
                f"{f(r, 'total_movements'):,.0f}",
                f"{f(r, 'otp15_pct'):.1f}%",
                r["representation_flag"],
            ]
            for r in countries_best[:3] + countries_worst[:3]
        ],
        [2600, 1200, 1800, 1600, 3005],
    )
    add_callout(
        doc,
        "Jurisdiction interpretation",
        "The source uses ISO country-or-territory codes. Guernsey is a self-governing Crown Dependency, not a sovereign country. Aland Islands is also retained as an ISO-coded autonomous territory and leads the formal ranking with only one airport and 1,096 movements, so it is a limited screening result. Norway is the strongest broad, high-volume result at 94.9% OTP15. Canada is the weakest jurisdiction in the displayed bottom group for this selected international network population, not for all domestic Canadian aviation.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    section_break(doc, "4. Operating carriers and routes")
    add_body(
        doc,
        "AC Operator is the three-letter operating-carrier code, not necessarily the marketing airline shown to passengers. Unknown code ZZZ is excluded from named rankings.",
    )
    add_figure(
        doc,
        EXTRA / "figures" / "carrier_reliability_rankings.png",
        "Figure 8. Descriptive operating-carrier OTP15 with 95% Wilson intervals.",
        "OTP15 uses arrival delay. Eligibility: at least 1,000 flights and three periods; raw rankings are not adjusted for network mix.",
    )
    add_data_table(
        doc,
        ["Operating carrier", "Flights", "OTP15", "Wilson bound", "Screen"],
        [
            [f"{r['operator_name']} ({r['AC Operator']})", f"{f(r, 'flights'):,.0f}", f"{f(r, 'arrival_otp15_pct'):.1f}%", f"{f(r, 'arrival_otp15_ci_low_pct'):.1f}%", "Top 3"]
            for r in carrier_best[:3]
        ]
        + [
            [f"{r['operator_name']} ({r['AC Operator']})", f"{f(r, 'flights'):,.0f}", f"{f(r, 'arrival_otp15_pct'):.1f}%", f"{f(r, 'arrival_otp15_ci_high_pct'):.1f}%", "Bottom 3"]
            for r in carrier_worst[:3]
        ],
        [3500, 1500, 1500, 1800, 1905],
    )
    add_callout(
        doc,
        "Fair airline comparison",
        "The agreed adjusted OTP15 model should compare operating carriers after controlling route, origin, destination, scheduled duration, hour and period. Only carriers above 5,000 flights should be displayed in that adjusted league table.",
        LIGHT_BLUE,
        ORANGE,
    )

    page_heading(doc, "Operating-carrier delay burden scaled by duration", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "carrier_relative_delay_rankings.png",
        "Figure 8B. Operating carriers ranked by positive arrival-delay minutes relative to scheduled flight minutes.",
        "Eligibility remains at least 1,000 observed arrivals and three periods. Lower is better. This is an arrival-delay severity ratio, not the percentage of flights delayed.",
        16.4,
    )
    add_data_table(
        doc,
        ["Screen", "Operating carrier", "Flights", "OTP15", "Relative burden"],
        [
            [
                "Lowest" if row in carrier_duration_best[:3] else "Highest",
                f"{row['operator_name']} ({row['AC Operator']})",
                f"{f(row, 'observed_arrivals'):,.0f}",
                f"{f(row, 'arrival_otp15_pct'):.1f}%",
                f"{f(row, 'relative_delay_burden_pct'):.1f}%",
            ]
            for row in carrier_duration_best[:3] + carrier_duration_worst[:3]
        ],
        [1200, 3900, 1500, 1500, 2105],
    )
    add_callout(
        doc,
        "Why the ranking changes",
        "Long scheduled flights dilute the same number of delay minutes more than short flights. Therefore a low relative burden does not automatically mean high OTP15. The chart includes all regular operating carriers, including cargo operators; use it as a severity screen rather than a passenger-airline league table.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    page_heading(doc, "EUROCONTROL use case: Ryanair versus Wizz Air", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "ryanair_wizz_shared_routes.png",
        "Figure 8C. Ryanair and Wizz Air compared on their highest-volume shared directional routes.",
        "Brand groups include RYR/RYS/RUK and WZZ/WUK/WMT/WAZ. Every eligible shared route has at least 20 observed arrivals and three periods for each group.",
        16.7,
    )
    add_data_table(
        doc,
        ["Airline group", "Shared routes", "Balanced weight", "Std. OTP15", "Relative burden"],
        [
            [
                row["airline_family"],
                f"{f(row, 'shared_directional_routes'):,.0f}",
                f"{f(row, 'balanced_route_weight'):,.0f}",
                f"{f(row, 'route_standardised_otp15_pct'):.1f}%",
                f"{f(row, 'route_standardised_relative_delay_burden_pct'):.1f}%",
            ]
            for row in ryanair_wizz
        ],
        [2100, 1600, 1800, 2650, 2055],
    )
    ryanair_row = next(row for row in ryanair_wizz if row["airline_family"] == "Ryanair Group")
    wizz_row = next(row for row in ryanair_wizz if row["airline_family"] == "Wizz Air Group")
    add_body(
        doc,
        f"Across {f(ryanair_row, 'shared_directional_routes'):.0f} eligible shared directional routes, the balanced route screen gives Ryanair {f(ryanair_row, 'route_standardised_otp15_pct'):.1f}% OTP15 and {f(ryanair_row, 'route_standardised_relative_delay_burden_pct'):.1f}% relative burden, versus {f(wizz_row, 'route_standardised_otp15_pct'):.1f}% and {f(wizz_row, 'route_standardised_relative_delay_burden_pct'):.1f}% for Wizz Air. The difference is small and descriptive: matching routes reduces network-mix bias but does not control hour, date, aircraft, airport congestion or disruption.",
    )

    page_heading(doc, "EUROCONTROL use case: Madrid-Milan over time", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "madrid_milan_case_study.png",
        "Figure 8D. Arrival OTP15 for the Madrid-Milan passenger-airline corridor by snapshot, carrier and scheduled departure hour.",
        "Madrid is LEMD; the Milan airport system is LIMC, LIML and LIME; both directions are included. Passenger airlines are identified through mainline, low-cost or regional market segments.",
        17.0,
    )
    madrid_best_period = max(madrid_periods, key=lambda row: f(row, "arrival_otp15_pct"))
    madrid_worst_period = min(madrid_periods, key=lambda row: f(row, "arrival_otp15_pct"))
    madrid_best_operator = max(madrid_operators, key=lambda row: f(row, "arrival_otp15_pct"))
    madrid_worst_operator = min(madrid_operators, key=lambda row: f(row, "arrival_otp15_pct"))
    madrid_best_hour = max(madrid_hours, key=lambda row: f(row, "arrival_otp15_pct"))
    madrid_worst_hour = min(madrid_hours, key=lambda row: f(row, "arrival_otp15_pct"))
    add_callout(
        doc,
        "Corridor screen",
        f"The strongest snapshot is {madrid_best_period['period']} at {f(madrid_best_period, 'arrival_otp15_pct'):.1f}% OTP15; the weakest is {madrid_worst_period['period']} at {f(madrid_worst_period, 'arrival_otp15_pct'):.1f}%. Among airlines with at least 50 observed arrivals, {madrid_best_operator['operator_name']} ({madrid_best_operator['AC Operator']}) is highest at {f(madrid_best_operator, 'arrival_otp15_pct'):.1f}% and {madrid_worst_operator['operator_name']} ({madrid_worst_operator['AC Operator']}) lowest at {f(madrid_worst_operator, 'arrival_otp15_pct'):.1f}%.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    page_heading(doc, "Madrid-Milan: decision windows and individual records", level=2)
    add_data_table(
        doc,
        ["Decision screen", "Best", "OTP15", "Worst", "OTP15"],
        [
            [
                "Snapshot",
                madrid_best_period["period"],
                f"{f(madrid_best_period, 'arrival_otp15_pct'):.1f}%",
                madrid_worst_period["period"],
                f"{f(madrid_worst_period, 'arrival_otp15_pct'):.1f}%",
            ],
            [
                "Passenger airline",
                f"{madrid_best_operator['AC Operator']} ({f(madrid_best_operator, 'observed_arrivals'):,.0f} flights)",
                f"{f(madrid_best_operator, 'arrival_otp15_pct'):.1f}%",
                f"{madrid_worst_operator['AC Operator']} ({f(madrid_worst_operator, 'observed_arrivals'):,.0f} flights)",
                f"{f(madrid_worst_operator, 'arrival_otp15_pct'):.1f}%",
            ],
            [
                "Scheduled departure hour",
                f"{int(f(madrid_best_hour, 'departure_hour')):02d}:00 ({f(madrid_best_hour, 'observed_arrivals'):,.0f} flights)",
                f"{f(madrid_best_hour, 'arrival_otp15_pct'):.1f}%",
                f"{int(f(madrid_worst_hour, 'departure_hour')):02d}:00 ({f(madrid_worst_hour, 'observed_arrivals'):,.0f} flights)",
                f"{f(madrid_worst_hour, 'arrival_otp15_pct'):.1f}%",
            ],
        ],
        [1850, 2700, 1500, 2700, 1455],
    )
    doc.add_heading("Largest observed arrival delays in the corridor", level=3)
    add_data_table(
        doc,
        ["Date", "Route", "Operator", "Scheduled duration", "Arrival delay"],
        [
            [
                row["departure_date"],
                row["route"],
                row["AC Operator"],
                f"{f(row, 'scheduled_duration_min'):.0f} min",
                f"{f(row, 'Arrival_Delay_Min'):.0f} min",
            ]
            for row in madrid_extremes[:3]
        ],
        [1900, 2200, 1500, 2300, 2305],
    )
    add_body(
        doc,
        "This is how EUROCONTROL data supports action: identify a corridor deterioration, locate the affected months and hours, compare the operators serving the same market, and then retrieve individual records for operational audit. The records are examples, not proof that the operator caused the delay; weather, reactionary delay, airport constraints and schedule definitions are not yet controlled.",
    )

    page_heading(doc, "Highest-volume route context", level=2)
    add_figure(
        doc,
        EDA / "figures" / "top_route_comparison.png",
        "Figure 9. Flight volume and delay-above-15 rate on the highest-volume eligible directional routes.",
        "Executive route eligibility is at least 20 flights and three periods; the display is restricted to the largest routes.",
    )

    page_heading(doc, "Executive exceptions: time, routes and carriers", level=2)
    add_body(
        doc,
        "This page restores the operational exceptions hidden by the high-volume overview. Day-by-hour cells use observed arrivals only; route and carrier results remain descriptive screens, not causal performance scores.",
    )
    best_cells = _rank(day_hours, "otp15_pct", descending=True, n=3)
    worst_cells = _rank(day_hours, "otp15_pct", n=3)
    add_data_table(
        doc,
        ["Screen", "Scheduled departure window", "Flights", "OTP15", "Wilson 95%"],
        [
            [
                "Best" if row in best_cells else "Worst",
                f"{row['departure_weekday']} {int(f(row, 'departure_hour')):02d}:00",
                f"{f(row, 'flights'):,.0f}",
                f"{f(row, 'otp15_pct'):.1f}%",
                f"{f(row, 'otp15_ci_low_pct'):.1f}–{f(row, 'otp15_ci_high_pct'):.1f}%",
            ]
            for row in best_cells + worst_cells
        ],
        [1200, 2850, 1700, 1450, 3005],
    )
    doc.add_heading("Five directional routes requiring data and schedule review", level=3)
    add_data_table(
        doc,
        ["Route", "Flights", "Periods", "OTP15", "Wilson 95%"],
        [
            [
                f"{row['ADEP']} → {row['ADES']}",
                f"{f(row, 'flights'):,.0f}",
                f"{f(row, 'periods_active'):.0f}",
                f"{f(row, 'arrival_otp15_pct'):.1f}%",
                f"{f(row, 'arrival_otp15_ci_low_pct'):.1f}–{f(row, 'arrival_otp15_ci_high_pct'):.1f}%",
            ]
            for row in problematic_routes[:5]
        ],
        [2100, 1600, 1300, 1500, 3705],
    )
    add_body(
        doc,
        "Long-haul warning. These extreme directional routes may expose schedule-time, timezone, date-rollover or coverage problems. Audit those definitions before treating the result as operational underperformance.",
    )
    doc.add_heading("Operating-carrier comparison status", level=3)
    add_data_table(
        doc,
        ["Raw bottom screen", "Flights", "OTP15", "Adjusted result"],
        [
            [
                f"{row['operator_name']} ({row['AC Operator']})",
                f"{f(row, 'flights'):,.0f}",
                f"{f(row, 'arrival_otp15_pct'):.1f}%",
                "Pending model execution",
            ]
            for row in carrier_worst[:3]
        ],
        [3500, 1500, 1500, 3705],
    )
    add_body(
        doc,
        "Adjusted carrier comparison is not yet available. The agreed model will control route, origin, destination, scheduled duration, departure hour and period, displaying only operators above 5,000 flights.",
    )

    section_break(doc, "5. Congestion")
    add_body(
        doc,
        "Congestion is not directly observed. The proxy ranks each airport-date-hour against that origin airport's own traffic distribution and assigns a low, moderate, high or peak quartile. This compares each airport with itself.",
    )
    add_figure(
        doc,
        EXTRA / "figures" / "congestion_pressure_performance.png",
        "Figure 10. Delay rate across within-airport traffic-load quartiles.",
        "Association only: pressure is scheduled operated departures within each origin airport.",
    )
    low, peak = pressure[0], pressure[-1]
    add_callout(
        doc,
        "Observed pressure gradient",
        f"The delay-above-15 rate rises from {f(low, 'delayed15_pct'):.1f}% in low-load airport-hours to {f(peak, 'delayed15_pct'):.1f}% in peak-load airport-hours, a {f(peak, 'delayed15_pct') - f(low, 'delayed15_pct'):.1f}-point difference.",
        LIGHT_BLUE,
        DARK_BLUE,
    )
    add_body(
        doc,
        "The result can guide staffing, stand allocation and monitoring, but does not establish that extra traffic caused the delay. Peak periods can also differ in route mix, weather and accumulated rotation delay.",
    )

    page_heading(doc, "Delay rates by flight duration", level=2)
    add_figure(
        doc,
        EXTRA / "figures" / "duration_delay_context.png",
        "Figure 11. Arrival-delay severity and traffic exposure by scheduled flight-duration band.",
        "Denominator: observed arrivals with a positive scheduled duration. Delay rates show arrivals more than 15, 30 and 60 minutes late; the right panel shows how much traffic each band represents.",
        17.2,
    )
    duration_rows = rows(tables / "duration_delay_performance.csv")
    shortest = duration_rows[0]
    longest = duration_rows[-1]
    add_body(
        doc,
        f"Flights of 90 minutes or less account for {f(shortest, 'flight_share_pct'):.1f}% of the analysed traffic and have a {f(shortest, 'delayed_15_pct'):.1f}% delay-above-15 rate. Flights longer than six hours account for {f(longest, 'flight_share_pct'):.1f}% and have a {f(longest, 'delayed_15_pct'):.1f}% rate. The 30- and 60-minute bars distinguish routine lateness from more severe disruption.",
    )
    add_callout(
        doc,
        "Business interpretation",
        "Duration bands add context to raw airport, route and carrier OTP15 comparisons. The gradient is descriptive, not a causal duration effect; the unusually high long-haul rate also reinforces the need to audit schedule times, timezones and date rollover. A group with a different route-length mix should not be judged from its unadjusted punctuality rate alone.",
        LIGHT_BLUE,
        DARK_BLUE,
    )

    section_break(doc, "6. Weekday and hour")
    add_figure(
        doc,
        EXTRA / "figures" / "weekday_otp15_performance.png",
        "Figure 12. Weekday OTP15 with 95% Wilson intervals.",
        "Each weekday is also tested against the rest of the network with Benjamini-Hochberg adjustment.",
    )
    add_figure(
        doc,
        EDA / "figures" / "time_reliability_heatmap.png",
        "Figure 13. OTP15 by scheduled departure weekday and hour.",
        "The matrix replaces separate hour-only charts that added little decision value.",
    )
    best_day = max(weekdays, key=lambda row: f(row, "otp15_pct"))
    worst_day = min(weekdays, key=lambda row: f(row, "otp15_pct"))
    add_body(
        doc,
        f"{best_day['departure_weekday']} is descriptively strongest at {f(best_day, 'otp15_pct'):.1f}% OTP15 and {worst_day['departure_weekday']} weakest at {f(worst_day, 'otp15_pct'):.1f}%. All weekday-versus-rest tests are significant, but none reaches the three-point practical threshold. The heatmap is retained because operational windows can be hidden by a daily average.",
    )

    section_break(doc, "7. Change monitoring")
    add_figure(
        doc,
        EXTRA / "figures" / "origin_airport_otp15_trend_extremes.png",
        "Figure 14. Three material origin-airport OTP15 deteriorations in the high-volume pool.",
        "At least six snapshots; the asterisk requires adjusted significance and an absolute change of at least 3 points.",
    )
    add_figure(
        doc,
        EXTRA / "figures" / "operator_otp15_trend_extremes.png",
        "Figure 15. Three material operating-carrier OTP15 deteriorations in the high-volume pool.",
        "Trend magnitude can reflect coverage, route mix or exceptional periods and must be audited before attribution.",
    )
    add_body(
        doc,
        "The displayed high-volume pools contain material deteriorations but no improvements that pass both the adjusted-significance and three-point rules. LPPT and TAP are the largest-volume entities among the displayed severe airport and carrier signals. Several estimated changes are unusually large, so the first business action is a coverage and schedule-time audit, not a performance accusation.",
    )

    section_break(doc, "8. Recommendations and limitations")
    doc.add_heading("Recommended investigations", level=2)
    add_body(doc, "Airport action. Investigate high-volume exceptions using matched route, operating carrier and hour slices. Treat long-haul origin extremes as a data-definition check before a local operational finding.")
    add_body(doc, "Carrier action. Fit the agreed adjusted OTP15 comparison for operators above 5,000 flights, controlling route, origin, destination, scheduled duration, hour and period. Publish adjusted effects and uncertainty alongside the raw screen.")
    add_body(doc, "Congestion action. Monitor the busiest airport-hour quartile and test whether staffing, stand planning or schedule smoothing changes OTP15 within the same airport.")
    add_body(doc, "Data action. Acquire consecutive monthly data, cancellations, diversions and passenger or seat exposure. Add weather after the flight-only baseline is frozen. Consecutive periods are more urgent than extra charts because they separate persistent change from isolated snapshots.")
    doc.add_heading("Limitations", level=2)
    add_body(doc, "The source covers operated flights and omits cancelled services. It contains no passenger counts, revenue, gate assignments or direct congestion measure. Monthly snapshots are non-consecutive. Airport, country-or-territory and carrier results are descriptive and can reflect route mix. P-values become very small at this scale, which is why the three-point rule is mandatory.")
    add_callout(
        doc,
        "Decision-ready conclusion",
        "The evidence is strong enough to prioritise high-volume airport, carrier and congestion investigations, but not to assign responsibility. Volume thresholds, Wilson intervals and the three-point rule stop statistical significance being mistaken for operational importance.",
        GREEN,
        DARK_BLUE,
    )

    doc.add_page_break()
    doc.add_heading("Appendix. Auditable outputs", level=1)
    products = [
        ("top200_airports_by_total_movements.csv", "Top-200 airport country, coordinates, movement volume and OTP15."),
        ("country_statistics_min_1000_movements.csv", "All country results above 1,000 historical movements."),
        ("weekday_vs_rest_otp15_tests.csv", "Weekday effect sizes, adjusted p-values and practical decisions."),
        ("top10_origin_airport_period_alerts.csv", "Leave-one-period-out alerts for the ten busiest origins."),
        ("origin_airport_otp15_trends.csv", "Snapshot trend tests for eligible origin airports."),
        ("operator_otp15_trends.csv", "Snapshot trend tests for eligible operating carriers."),
        ("operator_relative_delay_performance.csv", "Operating-carrier delay minutes scaled by scheduled flight duration."),
        ("ryanair_wizz_shared_directional_routes.csv", "Ryanair and Wizz Air results on eligible shared directional routes."),
        ("madrid_milan_period_performance.csv", "Madrid-Milan passenger-airline performance by snapshot."),
        ("madrid_milan_operator_performance.csv", "Madrid-Milan passenger-airline comparison."),
        ("madrid_milan_hour_performance.csv", "Madrid-Milan scheduled-departure-hour screen."),
    ]
    add_data_table(doc, ["CSV", "Purpose"], products, [3900, 6305])
    add_body(
        doc,
        "Dense exploratory charts and all calculation tables remain in the analysis outputs for traceability. They were omitted from this executive report when they repeated a clearer chart or relied on low-volume entities.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
