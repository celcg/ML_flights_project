"""Build the verified English-language business aviation report DOCX."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
EDA = ROOT / "reports" / "business_eda"
EXTRA = ROOT / "reports" / "business_report_full"
OUTPUT = ROOT / "doc" / "European_Scheduled_Aviation_Business_Report.docx"
BLUE = "5B9BD5"
LIGHT_BLUE = "D9EAF7"
GREEN = "70AD47"
DARK_BLUE = "1F4E78"
GREY = "667085"
LIGHT_GREY = "F2F4F7"
ORANGE = "ED7D31"
WHITE = "FFFFFF"
BLACK = "1F2937"
CONTENT_DXA = 10205  # A4 with 1.5 cm margins.


def rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def f(row, key):
    return float(row[key])


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=12, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 9, GREY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(.65)
    section.footer_distance = Cm(.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 18, DARK_BLUE, 16, 8),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("EUROPEAN SCHEDULED AVIATION PERFORMANCE  |  BUSINESS ANALYSIS")
    set_font(run, 8.5, GREY, bold=True)
    add_page_number(section.footer.paragraphs[0])


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("EUROPEAN SCHEDULED\nAVIATION PERFORMANCE"), 28, DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    set_font(p.add_run("Demand, punctuality, airport pressure and operational recovery"), 15, BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    set_font(p.add_run("Nine observed monthly periods | June 2021 to June 2023"), 11, GREY, italic=True)

    table = doc.add_table(rows=2, cols=3)
    set_table_geometry(table, [3402, 3402, 3401])
    cards = [
        ("5.38m", "in-scope operated flights"),
        ("82.2%", "arrival OTP15"),
        ("22.0 min", "p90 arrival delay"),
        ("34,420", "directional routes"),
        ("485", "operating-carrier codes"),
        ("9", "monthly snapshots"),
    ]
    for cell, (value, label) in zip([c for row in table.rows for c in row.cells], cards):
        set_cell_fill(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(value), 18, DARK_BLUE, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 9.5, GREY)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    set_font(p.add_run("Prepared from the complete business EDA execution"), 11, DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Descriptive evidence for operational prioritisation; not a causal or passenger-demand study"), 10, GREY)
    doc.add_page_break()


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "  "), 11, accent, bold=True)
    set_font(p.add_run(text), 11, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_font(p.add_run(bold_lead), 12, BLACK, bold=True)
        set_font(p.add_run(text[len(bold_lead):]), 12, BLACK)
    else:
        set_font(p.add_run(text), 12, BLACK)
    return p


def add_figure(doc, image, caption, source, width_cm=17.5):
    path = Path(image)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Cm(width_cm))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(2)
    set_font(cap.add_run(caption), 9, GREY, italic=True)
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.paragraph_format.space_before = Pt(0)
    src.paragraph_format.space_after = Pt(6)
    set_font(src.add_run(source), 8.5, GREY)


def add_data_table(doc, headers, data, widths, formats=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(header), 9, WHITE, bold=True)
    for i, record in enumerate(data):
        row = table.add_row()
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = row.cells
        for j, (cell, value) in enumerate(zip(cells, record)):
            if i % 2:
                set_cell_fill(cell, "F7FBFE")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(str(value)), 9, BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def section_break(doc, title):
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.page_break_before = True


def page_heading(doc, title, level=2):
    heading = doc.add_heading(title, level=level)
    heading.paragraph_format.page_break_before = True
    return heading


def build():
    kpi_rows = rows(EDA / "tables" / "network_kpis.csv")
    kpis = {row[""]: float(row["value"]) for row in kpi_rows}
    top_airports = rows(EXTRA / "tables" / "top200_airports_by_total_movements.csv")
    countries = rows(EXTRA / "tables" / "top200_airport_country_statistics.csv")
    carriers = rows(EXTRA / "tables" / "most_reliable_operating_carriers.csv")
    large_carriers = rows(EXTRA / "tables" / "largest_operating_carriers.csv")
    pressure = rows(EXTRA / "tables" / "within_airport_pressure_band_performance.csv")
    hours = rows(EXTRA / "tables" / "hourly_network_performance.csv")
    windows = rows(EXTRA / "tables" / "busiest_recurring_airport_hour_windows.csv")
    latlon = rows(EXTRA / "tables" / "latitude_longitude_delay_hotspots.csv")
    htests = rows(EDA / "statistical_tests" / "hypothesis_tests.csv")

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Executive summary", level=1)
    add_callout(doc, "Network result", "82.2% of the 5.38 million in-scope flights arrived no more than 15 minutes late. The median arrival delay was 2.2 minutes, but the p90 was 22.0 minutes, showing a material right tail.", GREEN, DARK_BLUE)
    add_body(doc, "The business value of this report is prioritisation: it identifies where disruption is concentrated by airport, country, operating carrier, traffic pressure and hour. It does not rank passenger experience because the source contains operated flights, not passengers, seats, cancellations or revenue.")
    add_body(doc, "Airport scale alone does not explain punctuality. EDDF and EGLL each recorded approximately 261 thousand endpoint movements, yet their combined OTP15 rates were 82.8% and 65.5% respectively. That 17.3-point gap is a stronger operational signal than volume by itself and warrants route-mix and process investigation.")
    add_body(doc, "Within each origin airport, peak-load airport-hours recorded 78.9% OTP15 versus 88.0% in low-load airport-hours. The 9.1-point difference is consistent with congestion pressure, although it is not causal proof: airport identity, route mix, weather and late-day propagation remain possible explanations.")
    add_body(doc, "Country aggregation of the Top 200 airports shows that Spain carried the largest represented volume (1.14 million endpoint movements) with 87.3% OTP15. The United Kingdom was second by volume (1.06 million) but recorded 79.1%, while Norway combined lower volume with the strongest scale-qualified country result at 94.2%.")
    add_callout(doc, "Priority actions", "Investigate EGLL and LFPG against same-scale peers; review recurring peak airport-hours; separate carrier performance from route mix with an adjusted model; acquire consecutive monthly observations before attributing changes to seasonality or policy.", LIGHT_BLUE, DARK_BLUE)

    section_break(doc, "1. Methodology and analytical decisions")
    doc.add_heading("Scope and denominator", level=2)
    add_body(doc, "The analysis uses nine EUROCONTROL monthly snapshots: June, September and December 2021; March, June, September and December 2022; and March and June 2023. These are observed snapshots rather than an uninterrupted 25-month time series.")
    add_body(doc, "Only regular scheduled flights were retained (ICAO Flight Type = S). Physical-quality rules excluded impossible early delays below -120 minutes, flight levels outside 0–500 and non-positive flown distance when those fields were observed. Null physical measurements were retained where the business KPI did not require them. Arrival analyses require an observed arrival time and use operated flights as the denominator.")
    add_body(doc, "Route graphics exclude routes with fewer than two historical flights. That rule prevents one-off records from entering charts, but it is not sufficient for executive ranking. Main route rankings therefore require at least 500 flights and three active periods; carrier rankings require at least 1,000 flights and three periods; airport executive tables require at least 1,000 flights and three periods. Country reliability rankings require at least three represented Top-200 airports and 20,000 endpoint movements.")
    doc.add_heading("Metrics and uncertainty", level=2)
    add_body(doc, "OTP15 is the share of flights arriving no more than 15 minutes late. Delay shares above 15, 30 and 60 minutes describe severity. The median describes a typical flight; p90 shows the threshold exceeded by the worst 10%, making the long disruption tail visible.")
    add_body(doc, "Wilson 95% confidence intervals accompany proportions because raw percentages can be misleading when sample sizes differ. Wilson remains bounded between 0% and 100% and behaves better than a simple normal interval for small samples or rates close to the extremes. Rankings use the conservative Wilson bound rather than the point estimate alone.")
    add_figure(doc, EDA / "figures" / "statistical_method_explainer.png",
               "Figure 1. Statistical methods used to separate observed rates, uncertainty and practical effect size.",
               "Source: complete EUROCONTROL business EDA; 5,375,605 in-scope flights.", 17.2)
    page_heading(doc, "Why p-tests were used", level=2)
    add_body(doc, "A p-test asks whether the observed data would be unusually incompatible with a stated null hypothesis. It does not measure business importance. The report therefore pairs every p-value with an effect size and applies Benjamini–Hochberg correction when multiple hypotheses are considered.")
    add_body(doc, "The two-proportion z-test compares December OTP15 rates; paired Wilcoxon tests whether delayed departures typically recover time; and Kruskal–Wallis compares delay distributions across duration bands without assuming normal delays. With millions of flights, tiny differences can be statistically significant, so effect sizes and operational thresholds drive interpretation.")
    add_callout(doc, "Denominator decision", "The source covers operated scheduled flights. It does not contain a complete cancellation or passenger denominator, so the report describes flight punctuality and operational exposure rather than total passenger disruption.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Ranking decision", "The two-flight rule removes one-off routes from charts. Executive rankings use much stricter volume and period requirements because a visually plausible percentage can still be unstable or unrepresentative.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Multiplicity decision", "Benjamini–Hochberg adjustment limits the expected false-discovery share when several hypotheses are tested. This prevents a long list of comparisons from being treated as independent evidence.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Complex decision", "All nine months are valid for descriptive reporting, but descriptive associations do not authorise post-event variables in the T−60 prediction model. Reporting scope and predictive leakage control are deliberately separate.", LIGHT_BLUE, DARK_BLUE)

    section_break(doc, "2. Airports and countries")
    doc.add_heading("Volume does not imply reliability", level=2)
    a1, a2 = top_airports[0], top_airports[1]
    add_body(doc, f"{a1['airport']} was the largest airport in the endpoint view with {f(a1,'total_movements'):,.0f} movements and {f(a1,'otp15_pct'):.1f}% OTP15. {a2['airport']} followed with {f(a2,'total_movements'):,.0f} movements but only {f(a2,'otp15_pct'):.1f}% OTP15. Similar scale can coexist with materially different reliability.")
    add_figure(doc, EDA / "figures" / "origin_airport_volume_reliability.png",
               "Figure 2. Origin-airport volume versus subsequent arrival reliability.",
               "Minimum 30 flights in the development chart; executive interpretations apply stricter volume checks.")
    add_figure(doc, EDA / "figures" / "destination_airport_volume_reliability.png",
               "Figure 3. Destination-airport volume versus arrival reliability.",
               "Minimum 30 flights; rates are descriptive and may reflect route and operator mix.")
    add_callout(doc, "Business interpretation", "EGLL is a priority exception: its 261,084 endpoint movements are virtually identical to EDDF's 261,184, but EGLL records 65.5% OTP15 versus EDDF's 82.8%. The scale-matched comparison suggests that capacity alone is not an adequate explanation.", LIGHT_BLUE, ORANGE)

    page_heading(doc, "Airport reliability rankings", level=2)
    add_figure(doc, EDA / "figures" / "origin_airport_reliability_rankings.png",
               "Figure 4. Most reliable and most problematic origin airports with 95% Wilson intervals.",
               "Interpretation concerns outcomes of flights departing each airport, not an adjusted airport causal effect.", 15.8)
    add_figure(doc, EDA / "figures" / "destination_airport_reliability_rankings.png",
               "Figure 5. Most reliable and most problematic destination airports with 95% Wilson intervals.",
               "The report's executive tables require at least 1,000 flights and three active periods.", 15.8)
    add_body(doc, "The exploratory chart's first three problematic origins are CYUL, KBNA and CYYT. However, KBNA and CYYT have only 144 and 58 observed flights, so they do not satisfy the executive eligibility rule. After requiring at least 1,000 flights and three active periods, the three most problematic origins are CYUL, CYYZ and CYVR. This should not be read as proof of poor local airport management: the selected population is flights interacting with the EUROCONTROL area and is strongly affected by route direction and schedule conventions. Destination results identify OMDB and EGLL as high-volume attention points, but still require adjusted follow-up.")

    page_heading(doc, "Geographic concentration and the Top 200 airports", level=2)
    add_figure(doc, EXTRA / "figures" / "top200_airport_volume_delay_map.png",
               "Figure 6. Coordinate map of the 200 largest airports; bubble area is volume and colour is the >15-minute delay share.",
               "Top 200 defined by departures plus arrivals. All 200 airports were matched to coordinates.")
    add_figure(doc, EXTRA / "figures" / "latitude_longitude_delay_hotspots.png",
               "Figure 7. Delayed movements accumulated in 5° latitude × 5° longitude bands.",
               "Metric is the count of movements delayed >15 minutes, preventing early arrivals from cancelling severe delays.")
    hotspot = latlon[0]
    add_callout(doc, "Largest geographic accumulation", f"The {hotspot['latitude_band']}°–{int(float(hotspot['latitude_band']))+5}° latitude and {hotspot['longitude_band']}°–{int(float(hotspot['longitude_band']))+5}° longitude band contains {f(hotspot,'delayed15_count'):,.0f} delayed endpoint movements across {f(hotspot,'movements'):,.0f} movements ({f(hotspot,'delayed15_pct'):.1f}%). This is an exposure hotspot, not necessarily the worst rate.", LIGHT_BLUE, DARK_BLUE)

    doc.add_heading("Country statistics from the Top 200", level=2)
    add_body(doc, "Countries are assigned from the ICAO airport dimension. The unit is an airport endpoint movement, so each flight may contribute once at its origin and once at its destination when both airports are in the Top 200. Results describe the represented airport network, not national passenger demand.")
    add_figure(doc, EXTRA / "figures" / "country_top200_delay_hotspot_map.png",
               "Figure 8. Country delay hotspots represented by their Top-200 airports.",
               "Bubble area is endpoint movements; colour is the share delayed >15 minutes; position is the mean airport coordinate.")
    add_figure(doc, EXTRA / "figures" / "country_top200_reliability_ranking.png",
               "Figure 9. Scale-qualified country reliability among the Top-200 airports.",
               "Eligibility: at least three represented airports and 20,000 endpoint movements.")
    top_country_rows = countries[:8]
    add_data_table(doc,
        ["Country", "Airports", "Movements", "OTP15", "Delay >15"],
        [[r["country"], int(f(r,"airports")), f"{f(r,'total_movements')/1e6:.2f}m",
          f"{f(r,'otp15_pct'):.1f}%", f"{f(r,'delayed15_pct'):.1f}%"] for r in top_country_rows],
        [3000, 1300, 1900, 1900, 2105])
    add_callout(doc, "Country comparison", "Spain combines the largest represented volume with 87.3% OTP15. The United Kingdom has only 7% fewer endpoint movements but an 8.2-point lower OTP15. Norway leads the scale-qualified country ranking at 94.2%, although its volume and airport mix differ substantially.", LIGHT_BLUE, DARK_BLUE)

    section_break(doc, "3. Operating carriers")
    add_body(doc, "AC Operator is the three-letter operating-carrier code, not necessarily the marketing airline shown to passengers. Unknown code ZZZ is excluded from named rankings. Company labels come from the ICAO dimension and may require periodic reference-data refresh.")
    add_figure(doc, EXTRA / "figures" / "carrier_reliability_rankings.png",
               "Figure 10. Descriptive operating-carrier reliability with 95% Wilson intervals.",
               "Eligibility: at least 1,000 flights and three periods; rankings are not adjusted for route mix.")
    add_data_table(doc,
        ["Rank", "Operating carrier", "Flights", "OTP15", "p90 delay", "Routes"],
        [[i+1, f"{r['operator_name']} ({r['AC Operator']})", f"{f(r,'flights'):,.0f}",
          f"{f(r,'arrival_otp15_pct'):.1f}%", f"{f(r,'arrival_delay_p90'):.1f} min", f"{f(r,'routes'):,.0f}"]
         for i, r in enumerate(carriers[:5])],
        [700, 3300, 1500, 1500, 1700, 1505])
    add_body(doc, "The formal Top 3 are JOON (JON), Widerøe (WIF) and Braathens Regional Airways (BRX), ranked by the lower Wilson bound. JOON has only 2,275 observed flights, whereas Widerøe has 88,523. For an operational scale comparison requiring at least 10,000 flights, the Top 3 become Widerøe, Braathens Regional Airways and Binter Canarias (IBB).")
    add_callout(doc, "Why no causal airline league table", "Carrier outcomes depend on routes, airports, departure times, duration and schedule buffers. The fair next comparison is an adjusted delay model with route and time controls; the present ranking is a screening tool, not a performance verdict.", LIGHT_BLUE, ORANGE)
    doc.add_heading("Scale and breadth", level=2)
    large = large_carriers[:10]
    add_data_table(doc,
        ["Carrier", "Flights", "Routes", "Airports", "OTP15"],
        [[f"{r['operator_name']} ({r['AC Operator']})", f"{f(r,'flights'):,.0f}", f"{f(r,'routes'):,.0f}",
          f"{f(r,'airports'):,.0f}", f"{f(r,'arrival_otp15_pct'):.1f}%"] for r in large],
        [3500, 1800, 1600, 1600, 1705])
    add_body(doc, "Ryanair is the largest identifiable operator with 634,990 flights and 86.9% OTP15. Lufthansa records 87.8% across 257,120 flights, while British Airways records 63.0% across 120,845. These gaps are large enough to prioritise investigation, but not to assign responsibility before adjusting for network mix.")

    section_break(doc, "4. Congestion and operational pressure")
    doc.add_heading("Route volume versus reliability", level=2)
    add_figure(doc, EDA / "figures" / "route_volume_reliability.png",
               "Figure 11. Directional-route volume versus OTP15.",
               "No route with fewer than two historical flights is plotted; executive route ranking uses at least 500 flights and three periods.")
    add_figure(doc, EDA / "figures" / "top_route_comparison.png",
               "Figure 12. Reliability comparison among the highest-volume directional routes.",
               "Percentages use operated scheduled flights with observed arrival outcome.")
    add_body(doc, "The route view distinguishes repeated operational exposure from isolated poor percentages. The threshold analysis shows that 2,800 routes meet the 500-flight and three-period executive rule and cover 55.5% of all analysed flights; using only the very largest routes would miss a substantial share of the network.")

    doc.add_heading("Within-airport traffic pressure", level=2)
    add_body(doc, "Congestion is not directly observed in the dataset. A transparent proxy was therefore built: for each origin airport, every airport-date-hour was ranked against that airport's own traffic distribution and assigned to a low, moderate, high or peak quartile. This compares an airport with itself rather than treating ten flights as equally congested everywhere.")
    add_figure(doc, EXTRA / "figures" / "congestion_pressure_performance.png",
               "Figure 13. Arrival delay rate by within-airport traffic-load quartile.",
               "Association only: pressure band is scheduled operated departures within each origin airport.")
    low, peak = pressure[0], pressure[-1]
    add_callout(doc, "Observed pressure gradient", f"The >15-minute delay rate rises from {f(low,'delayed15_pct'):.1f}% in low-load airport-hours to {f(peak,'delayed15_pct'):.1f}% in peak-load airport-hours. Mean arrival delay rises from {f(low,'arrival_delay_mean'):.1f} to {f(peak,'arrival_delay_mean'):.1f} minutes.", LIGHT_BLUE, DARK_BLUE)
    add_body(doc, "The result is operationally useful because it identifies high-load windows for staffing, stand allocation and disruption monitoring. It cannot establish that extra traffic caused the delay; peak periods may also differ in airport, route, weather and accumulated rotation delay.")
    doc.add_heading("Recurring high-volume windows", level=2)
    add_data_table(doc,
        ["Origin", "Hour", "Mean dep./day", "Active days", "OTP15", "Delay >15"],
        [[r["ADEP"], f"{int(f(r,'departure_hour')):02d}:00", f"{f(r,'mean_departures_per_active_day'):.1f}",
          int(f(r,"active_days")), f"{f(r,'otp15_pct'):.1f}%", f"{f(r,'delayed15_pct'):.1f}%"] for r in windows[:6]],
        [1300, 1300, 2000, 1700, 1900, 2005])
    add_body(doc, "EDDF at 11:00 is the busiest recurring airport-hour, averaging 44.0 departures per active day and 73.5% OTP15. LFPG at 08:00 and 10:00 combines slightly lower recurrent volume with only about 62% OTP15, making those windows more urgent monitoring candidates than volume alone suggests.")

    doc.add_heading("Delay propagation and recovery", level=2)
    add_figure(doc, EDA / "figures" / "departure_arrival_recovery.png",
               "Figure 14. Departure delay versus arrival delay and observed en-route recovery.",
               "Recovery is departure delay minus arrival delay; positive values indicate minutes recovered.", 16.0)
    add_body(doc, "Departure and arrival delays are strongly associated (Spearman rho 0.789), showing that late departure is the dominant retrospective signal. Among flights departing more than 15 minutes late, the median recovery is 1.8 minutes and 25.7% recover sufficiently to finish within OTP15. Recovery exists, but it does not usually erase major departure disruption.")

    section_break(doc, "5. Time-of-day performance")
    add_figure(doc, EDA / "figures" / "time_reliability_heatmap.png",
               "Figure 15. Arrival OTP15 by scheduled departure hour and weekday.",
               "Cell values are descriptive rates; local/UTC interpretation follows the source schedule convention.")
    add_figure(doc, EXTRA / "figures" / "hourly_network_volume_reliability.png",
               "Figure 16. Network flight volume and OTP15 by scheduled departure hour.",
               "All nine periods; denominator is operated scheduled flights with observed arrival outcome.")
    best = max(hours, key=lambda r: f(r, "otp15_pct"))
    worst = min(hours, key=lambda r: f(r, "otp15_pct"))
    add_callout(doc, "Hour comparison", f"The strongest network hour is {int(f(best,'departure_hour')):02d}:00 at {f(best,'otp15_pct'):.1f}% OTP15. The weakest is {int(f(worst,'departure_hour')):02d}:00 at {f(worst,'otp15_pct'):.1f}%, with a p90 arrival delay of {f(worst,'arrival_delay_p90'):.1f} minutes. Late-night results should be interpreted with route and airport mix, not as a clock-time causal effect.", LIGHT_BLUE, DARK_BLUE)
    add_body(doc, "Reliability is strongest around 06:00 and again in the 17:00–19:00 window in these snapshots. The sharp fall at 22:00–02:00 is operationally meaningful for overnight monitoring, but those hours contain fewer flights and a different long-haul mix. A follow-up adjusted analysis should control for origin, destination, operator and duration.")

    section_break(doc, "6. Statistical evidence and complex decisions")
    add_data_table(doc,
        ["Question", "Method", "Adjusted p", "Effect", "Business reading"],
        [
            ["December OTP15 equality", "Two-proportion z", "<0.001", "4.87 pp",
             "December 2021 exceeded December 2022; material change."],
            ["En-route recovery", "Paired Wilcoxon", "<0.001", "1.82 min median",
             "Recovery is real but modest for a typical delayed departure."],
            ["Equal delay by haul band", "Kruskal–Wallis", "<0.001", "epsilon² 0.056",
             "Distributions differ; duration band explains only part of variation."],
        ],
        [2450, 1800, 1300, 1700, 2955])
    add_body(doc, "The December test rejects equal OTP15 with a 4.87 percentage-point difference in favour of December 2021. This is both statistically and operationally material, but two December snapshots cannot identify why the change occurred.")
    add_body(doc, "The recovery test rejects a zero-median recovery effect, yet the median gain is only 1.82 minutes. The business conclusion is not that airlines can rely on recovery; it is that schedules contain some recoverable buffer, while departure prevention remains more valuable.")
    add_body(doc, "The haul-band test also rejects equality. Its epsilon-squared effect of 0.056 is meaningful but not dominant, supporting duration as a segmentation variable rather than a complete explanation of delay.")
    doc.add_heading("Key decisions that protect interpretation", level=2)
    add_callout(doc, "Volume thresholds", "A two-flight minimum protects plotting, while stricter 500-route / 1,000-airport or carrier thresholds protect executive ranking. Wilson bounds then account for remaining sample-size differences.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Country aggregation", "The Top 200 is defined by combined endpoint movements. National results are therefore exposure summaries of represented airports, not estimates of all national aviation or passenger performance.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Congestion proxy", "Traffic pressure is normalised within airport so each facility is compared with its own operating pattern. The result is a screening association, not a causal capacity estimate.", LIGHT_BLUE, DARK_BLUE)
    add_callout(doc, "Carrier fairness", "Raw carrier rates remain descriptive until route, airport, duration and time are controlled. This prevents a simple league table from becoming an unsupported performance verdict.", LIGHT_BLUE, DARK_BLUE)

    section_break(doc, "7. Recommendations and limitations")
    doc.add_heading("Recommended business investigations", level=2)
    add_body(doc, "1. Prioritise scale-matched airport exceptions. Compare EGLL with EDDF and LFPG with similarly large hubs, decomposing by route, hour, operator and turnaround pattern.")
    add_body(doc, "2. Monitor recurring high-load windows. Begin with LFPG 08:00/10:00, EDDF 11:00 and the EHAM morning/noon clusters; report both traffic and OTP15 weekly.")
    add_body(doc, "3. Replace raw carrier rankings with adjusted comparisons. Estimate operator effects after controlling for route, airport, scheduled duration and hour, then report odds ratios with uncertainty.")
    add_body(doc, "4. Acquire consecutive months. The current snapshots support broad prioritisation but cannot separate seasonality, persistent structural effects and temporary disruption.")
    add_body(doc, "5. Add cancellations, diversions, passenger/seat volume and gate/stand demand before converting flight reliability into customer or revenue impact. Weather remains a later enrichment after the flight-only expanded baseline is frozen.")
    doc.add_heading("Limitations", level=2)
    add_body(doc, "The dataset covers operated flights and therefore omits cancelled services. It contains no passenger counts, fares, revenue, gate assignments or direct congestion measure. Monthly snapshots are non-consecutive. Airport, country and carrier results are descriptive and may reflect route mix. P-values become very small with millions of records, so effect size and business context remain essential.")
    add_callout(doc, "Decision-ready conclusion", "The model-free evidence is strong enough to prioritise where to investigate, but not to assign responsibility. The most defensible next step is adjusted, time-aware analysis of high-volume airports, carriers and peak windows using consecutive periods.", GREEN, DARK_BLUE)

    doc.add_page_break()
    doc.add_heading("Appendix A. Delivered data products", level=1)
    add_body(doc, "The report is supported by auditable CSV outputs. The most important new files are listed below; complete airport, route and operator tables remain in reports/business_eda.")
    products = [
        ("top200_airports_by_total_movements.csv", "Top 200 airports with country, coordinates, volume, OTP15 and Wilson bounds."),
        ("top200_airport_country_statistics.csv", "Country aggregates over the Top-200 endpoint movements."),
        ("latitude_longitude_delay_hotspots.csv", "5° × 5° geographic delay accumulation grid."),
        ("most_reliable_operating_carriers.csv", "Wilson-ranked eligible operating carriers with ICAO company labels."),
        ("hourly_network_performance.csv", "Network volume, OTP15, mean, median and p90 by scheduled hour."),
        ("within_airport_pressure_band_performance.csv", "Traffic-pressure quartiles and arrival reliability."),
        ("busiest_recurring_airport_hour_windows.csv", "Recurring high-volume origin-airport hours."),
    ]
    add_data_table(doc, ["CSV", "Purpose"], products, [3900, 6305])
    add_body(doc, "All figures use the complete EDA population of 5,375,605 in-scope flights unless a Top-200 or eligibility rule is stated in the caption.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
