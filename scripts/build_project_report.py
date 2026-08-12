from __future__ import annotations

from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc"
ASSET_DIR = OUT_DIR / "_informe_assets"
OUTPUT = OUT_DIR / "Informe_proyecto_prediccion_retrasos_vuelos.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "2A9D8F"
ORANGE = "E69F00"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D8DEE6"
DARK_GRAY = "3F4A56"
WHITE = "FFFFFF"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color: str, size: int = 18) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_fixed_width(table, total_twips: int = 9360) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")


def add_field(run, instruction: str) -> None:
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string("68727D")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    # Captions are placed after their figure/table; keeping them with the next
    # paragraph would detach them from the visual they describe.
    caption.paragraph_format.keep_with_next = False


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("PREDICCIÓN DE RETRASO DE LLEGADA · T−60")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Informe de decisiones y resultados  ·  ")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("68727D")
    page_run = p.add_run()
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = RGBColor.from_string("68727D")
    add_field(page_run, "PAGE")


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        lead.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc: Document, text: str, *, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc: Document, title: str, text: str, *, color: str = BLUE, fill: str = LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, color)
    r = p.add_run(f"{title}. ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    return p


def add_table(doc: Document, headers, rows, widths, *, font_size=9.1, first_col_bold=False):
    assert sum(widths) == 9360, sum(widths)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for idx, (cell, header, width) in enumerate(zip(hdr.cells, headers, widths)):
        set_cell_width(cell, width)
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(str(header))
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    for row_idx, row_values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for idx, (cell, value, width) in enumerate(zip(row.cells, row_values, widths)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(str(value))
            r.font.name = "Calibri"
            r.font.size = Pt(font_size)
            if first_col_bold and idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    # Short comparison tables should move as a unit instead of leaving a
    # repeated header and a single orphan row on the previous page.
    if len(rows) <= 7:
        for table_row in table.rows[:-1]:
            for cell in table_row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc: Document, text: str):
    return doc.add_paragraph(text, style="Caption")


def add_page_break(doc: Document) -> None:
    # Keep the editorial cover on its own page. Later chapters flow naturally;
    # forced chapter breaks created nearly empty continuation pages in Word.
    count = getattr(add_page_break, "_count", 0)
    if count == 0:
        doc.add_page_break()
    add_page_break._count = count + 1


def add_small_label(doc: Document, text: str, color: str = TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def _font(size: int, bold: bool = False):
    font_name = "calibrib.ttf" if bold else "calibri.ttf"
    path = Path("C:/Windows/Fonts") / font_name
    if not path.exists():
        path = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    return ImageFont.truetype(str(path), size=size)


def _centered(draw, xy, text, font, fill):
    box = draw.textbbox((0, 0), text, font=font)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=font, fill=fill)


def create_assets() -> tuple[Path, Path, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    white = "#FFFFFF"
    grid = "#D8DEE6"
    text_color = "#3F4A56"
    title_font = _font(32, True)
    label_font = _font(23)
    small_font = _font(20)
    value_font = _font(20, True)

    # Dataset periods chart.
    period_path = ASSET_DIR / "periodos_dataset.png"
    img = Image.new("RGB", (1440, 594), white)
    d = ImageDraw.Draw(img)
    _centered(d, (720, 42), "Ampliación temporal del conjunto de vuelos", title_font, text_color)
    left, top, right, bottom = 115, 100, 1390, 500
    periods = ["dic-21", "mar-22", "jun-22", "sep-22", "dic-22", "mar-23"]
    counts = [0.5702, 0.576258, 0.813452, 0.826994, 0.64684, 0.681919]
    colors = ["#8BB8D8", "#71A8D0", "#5798C8", "#2E74B5", "#4CAFA7", "#2A9D8F"]
    for tick in (0.0, 0.25, 0.5, 0.75, 1.0):
        y = bottom - tick * (bottom - top)
        d.line((left, y, right, y), fill=grid, width=2)
        d.text((40, y - 13), f"{tick:.2f}", font=small_font, fill=text_color)
    slot = (right - left) / len(periods)
    bar_w = 112
    for i, (period, count, color) in enumerate(zip(periods, counts, colors)):
        cx = left + slot * (i + 0.5)
        y = bottom - count * (bottom - top)
        d.rounded_rectangle((cx - bar_w/2, y, cx + bar_w/2, bottom), radius=6, fill=color)
        _centered(d, (cx, y - 24), f"{count:.3f}", value_font, text_color)
        _centered(d, (cx, 538), period, label_font, text_color)
    d.text((8, 255), "Millones", font=small_font, fill=text_color)
    img.save(period_path, dpi=(180, 180))

    # Model comparison grouped bars.
    model_path = ASSET_DIR / "comparacion_mae.png"
    img = Image.new("RGB", (1476, 756), white)
    d = ImageDraw.Draw(img)
    _centered(d, (738, 42), "Evolución de los errores principales en validación", title_font, text_color)
    left, top, right, bottom = 105, 120, 1435, 610
    labels = ["Mediana\nglobal", "Ruta +\naerolínea", "Lineal\nregularizada", "CatBoost\n10%", "Ridge T−60\n10%", "CatBoost T−60\n10%"]
    global_mae = [11.685, 9.938, 10.871, 9.984, 9.941, 9.537]
    delayed_mae = [27.341, 21.109, 19.268, 20.484, 18.207, 19.635]
    max_y = 31.0
    for tick in (0, 10, 20, 30):
        y = bottom - tick / max_y * (bottom - top)
        d.line((left, y, right, y), fill=grid, width=2)
        d.text((55, y - 13), str(tick), font=small_font, fill=text_color)
    d.rounded_rectangle((955, 72, 980, 97), radius=4, fill="#2E74B5")
    d.text((990, 71), "MAE global", font=small_font, fill=text_color)
    d.rounded_rectangle((1180, 72, 1205, 97), radius=4, fill="#E69F00")
    d.text((1215, 71), "MAE >15 min", font=small_font, fill=text_color)
    slot = (right - left) / len(labels)
    bar_w = 54
    for i, label in enumerate(labels):
        cx = left + slot * (i + 0.5)
        for value, offset, color in ((global_mae[i], -30, "#2E74B5"), (delayed_mae[i], 30, "#E69F00")):
            y = bottom - value / max_y * (bottom - top)
            d.rounded_rectangle((cx + offset - bar_w/2, y, cx + offset + bar_w/2, bottom), radius=4, fill=color)
            _centered(d, (cx + offset, y - 18), f"{value:.1f}", _font(17, True), text_color)
        lines = label.split("\n")
        for line_i, line in enumerate(lines):
            _centered(d, (cx, 652 + line_i * 25), line, _font(18), text_color)
    d.text((8, 295), "Minutos", font=small_font, fill=text_color)
    img.save(model_path, dpi=(180, 180))

    # Operational-feature ablation horizontal chart.
    ablation_path = ASSET_DIR / "ablacion_variables.png"
    img = Image.new("RGB", (1440, 648), white)
    d = ImageDraw.Draw(img)
    _centered(d, (720, 42), "Las variables operacionales aportan más que la complejidad", title_font, text_color)
    labels = ["Estáticas", "+ aeropuertos", "+ ruta", "+ operador", "+ rotación", "Sin ventana 6 h"]
    scores = [15.612, 14.468, 14.318, 14.175, 14.153, 14.140]
    colors = ["#A9B3BF", "#8BB8D8", "#71A8D0", "#5798C8", "#3D8FBF", "#2A9D8F"]
    left, right, top = 330, 1340, 105
    x_min, x_max = 13.8, 15.9
    for tick in (14.0, 14.5, 15.0, 15.5):
        x = left + (tick - x_min) / (x_max - x_min) * (right - left)
        d.line((x, top, x, 555), fill=grid, width=2)
        _centered(d, (x, 584), f"{tick:.1f}", small_font, text_color)
    for i, (label, score, color) in enumerate(zip(labels, scores, colors)):
        y = top + i * 73
        _centered(d, (165, y + 25), label, label_font, text_color)
        x = left + (score - x_min) / (x_max - x_min) * (right - left)
        d.rounded_rectangle((left, y, x, y + 48), radius=6, fill=color)
        d.text((x + 12, y + 11), f"{score:.3f}", font=value_font, fill=text_color)
    _centered(d, (835, 625), "Puntuación combinada MAE (menor es mejor)", small_font, text_color)
    img.save(ablation_path, dpi=(180, 180))
    return period_path, model_path, ablation_path


def build_document() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    period_chart, model_chart, ablation_chart = create_assets()

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Informe del proyecto de predicción de retrasos de vuelos"
    doc.core_properties.subject = "Decisiones, resultados, conclusiones y próximos pasos"
    doc.core_properties.author = "Proyecto ML Flights"
    doc.core_properties.keywords = "vuelos, retraso de llegada, PySpark, Ridge, CatBoost, T-60"

    # Cover — editorial_cover template.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(78)
    spacer.paragraph_format.space_after = Pt(0)
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(14)
    r = label.add_run("INFORME DE DECISIONES Y RESULTADOS")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    r = title.add_run("Predicción del retraso\nde llegada de vuelos")
    r.font.name = "Calibri"
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(42)
    r = subtitle.add_run("Una hora antes de la salida · alcance actual arrival_pre")
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(DARK_GRAY)

    result = doc.add_paragraph()
    result.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result.paragraph_format.left_indent = Inches(0.55)
    result.paragraph_format.right_indent = Inches(0.55)
    result.paragraph_format.space_before = Pt(8)
    result.paragraph_format.space_after = Pt(30)
    result.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(result, LIGHT_BLUE)
    rr = result.add_run("Resultado actual: ")
    rr.bold = True
    rr.font.color.rgb = RGBColor.from_string(BLUE)
    result.add_run("Ridge con variables operacionales observables a T−60 logra un MAE global de 9,94 min y la mejor puntuación equilibrada (14,07) en validación.")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(42)
    meta.paragraph_format.space_after = Pt(0)
    r = meta.add_run("Versión 1.0  ·  9 de agosto de 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("68727D")

    add_page_break(doc)

    # Executive summary.
    add_small_label(doc, "Resumen ejecutivo")
    doc.add_heading("Qué se ha decidido y qué sabemos ya", level=1)
    add_body(doc, "El proyecto ha quedado centrado en una pregunta operativa concreta: estimar el retraso de llegada exactamente una hora antes de la salida programada. En ese momento no se puede usar el retraso de salida del propio vuelo, porque todavía no ha ocurrido. Sí se pueden usar datos planificados y resultados de vuelos anteriores que ya eran observables antes de ese corte.")
    add_callout(doc, "Conclusión principal", "La mejora más clara aparece al construir buenas variables temporales y operacionales, no al elegir un algoritmo cada vez más complejo. Ridge es por ahora el mejor modelo equilibrado; CatBoost es el mejor en error global y sigue siendo el principal competidor.")

    add_table(
        doc,
        ["Decisión", "Motivo", "Estado o consecuencia"],
        [
            ["Ampliar EUROCONTROL de un mes a seis periodos", "Un solo mes no representa estaciones ni cambios de tráfico.", "4.115.663 vuelos brutos entre dic-2021 y mar-2023."],
            ["Predecir solo llegada a T−60", "Es la necesidad prioritaria y evita mezclar horizontes operativos distintos.", "Variable objetivo: Arrival_Delay_Min."],
            ["Aplazar meteorología", "Antes conviene tener un baseline aeronáutico sólido y una unión temporal sin fugas.", "No se descarta; queda para una fase posterior."],
            ["Filtrar tráfico regular programado", "El modelo debe responder a una población homogénea y útil.", "3.671.258 vuelos limpios para modelado."],
            ["División temporal", "El futuro debe simular datos nuevos, no mezclarse al azar con el pasado.", "Train hasta sep-2022; validación dic-2022; test mar-2023."],
            ["Ridge como candidato actual", "Mejor equilibrio entre vuelos normales y retrasados, bajo consumo y fácil explicación.", "MAE global 9,94; MAE retrasados 18,21; combinado 14,07."],
        ],
        [2250, 3400, 3710],
        font_size=8.8,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 1. Resumen de las decisiones principales y su efecto en el proyecto.")

    doc.add_heading("Lectura rápida de los resultados", level=2)
    add_bullet(doc, "El baseline histórico de ruta + aerolínea supera con claridad a predecir siempre la mediana global.")
    add_bullet(doc, "Los primeros árboles (Random Forest, GBT y XGBoost) no mejoraron el equilibrio de la regresión lineal regularizada con el presupuesto pequeño de entrenamiento.")
    add_bullet(doc, "Log y Yeo–Johnson se probaron como transformaciones opcionales, aprendidas solo con train, pero no mejoraron la versión numérica original.")
    add_bullet(doc, "CatBoost mejora al pasar del 1% al 10% de entrenamiento, aunque las ganancias se reducen; más datos ayudarán, pero hacen falta mejores señales para dar un salto mayor.")
    add_bullet(doc, "Las ventanas de 1 y 24 horas, más la rotación del avión, mejoran especialmente los vuelos retrasados. La ventana de 6 horas se eliminó por no aportar valor adicional.")

    add_page_break(doc)

    # Scope.
    add_small_label(doc, "1 · Definición del problema")
    doc.add_heading("Una tarea concreta: retraso de llegada antes de salir", level=1)
    add_body(doc, "El objetivo actual es predecir Arrival_Delay_Min, es decir, la diferencia en minutos entre la hora real y la hora planificada de llegada. Un valor positivo indica retraso y uno negativo, llegada anticipada. La predicción se realiza en T−60: sesenta minutos antes de la hora programada de fuera de calzos (off-block).")

    doc.add_heading("Por qué fijar el horizonte T−60", level=2)
    add_body(doc, "Una predicción solo es útil si se puede calcular con la información disponible en el momento de tomar la decisión. Mezclar datos posteriores produce leakage o fuga de información: el modelo parece muy bueno durante el desarrollo, pero no podría repetir ese resultado en producción.")
    add_table(
        doc,
        ["Disponible a T−60", "No disponible para el propio vuelo"],
        [
            ["Aeropuerto de salida y llegada; operador; tipo de avión; segmento; nivel solicitado; horario y duración programada.", "Hora real de off-block; retraso de salida; hora real de llegada; distancia real volada."],
            ["Resultados de vuelos anteriores cuyo off-block o llegada ya sucedieron antes del corte.", "Información meteorológica futura o cualquier evento ocurrido después del corte."],
        ],
        [4680, 4680],
        font_size=9.2,
    )
    add_caption(doc, "Tabla 2. Contrato temporal usado para evitar fuga de información.")

    add_callout(doc, "Decisión descartada por ahora", "Usar el retraso de salida del mismo vuelo sirve para actualizar la predicción después del off-block, pero responde a otra pregunta y tendría ventaja informativa. Puede mantenerse como futuro modelo operacional post-off-block, nunca como comparación directa del modelo T−60.", color=ORANGE, fill="FFF5E6")

    doc.add_heading("Población incluida", level=2)
    add_body(doc, "El modelo se limita a vuelos regulares programados identificados por ICAO Flight Type = S. Tras aplicar este filtro, la variable queda constante y se elimina del modelo: sirve para definir el alcance, pero ya no puede explicar diferencias entre vuelos. Esta decisión también evita entrenar a la vez sobre operaciones regulares y no regulares, que pueden tener comportamientos distintos.")

    doc.add_heading("Configuración central del proyecto", level=2)
    add_body(doc, "Las reglas compartidas están concentradas en src/flight_config.py. Ahí se fijan el objetivo, el horizonte, los límites de calidad, el umbral de categorías raras y los interruptores opcionales de log y Yeo–Johnson. Centralizar estas decisiones evita que cada libreta aplique una versión diferente del problema.")

    add_page_break(doc)

    # Data.
    add_small_label(doc, "2 · Datos")
    doc.add_heading("Por qué se eligió y se amplió el dataset", level=1)
    add_body(doc, "La fuente principal es EUROCONTROL ADR. Se eligió porque ofrece millones de movimientos con una estructura coherente y campos directamente relacionados con la operación: horarios planificados y reales, aeropuertos, operador, tipo de aeronave, segmento y nivel de vuelo solicitado. Es una base adecuada para construir primero un modelo aeronáutico sin depender de fuentes externas.")

    doc.add_heading("De un mes a seis periodos", level=2)
    add_body(doc, "El primer análisis se hizo con diciembre de 2021: 570.200 vuelos. Era suficiente para entender columnas, nulos y distribuciones, pero no para confiar en la generalización. Por eso no se cambió a una fuente distinta: se añadieron más periodos de la misma fuente, espaciados para cubrir estaciones y cambios de volumen.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(period_chart), width=Inches(6.45))
    add_caption(doc, "Figura 1. Seis periodos mensuales suman 4.115.663 vuelos. Son cortes estacionales, no un calendario continuo.")

    add_table(
        doc,
        ["Periodo", "Vuelos", "Uso en el diseño temporal"],
        [
            ["dic-2021", "570.200", "Entrenamiento"],
            ["mar-2022", "576.258", "Entrenamiento"],
            ["jun-2022", "813.452", "Entrenamiento"],
            ["sep-2022", "826.994", "Entrenamiento"],
            ["dic-2022", "646.840", "Validación"],
            ["mar-2023", "681.919", "Test"],
        ],
        [1800, 1800, 5760],
        font_size=9,
    )

    doc.add_heading("Ventaja y límite de esta ampliación", level=2)
    add_body(doc, "La ventaja es que el modelo ve invierno, primavera, verano y otoño con un volumen aún manejable en un ordenador personal. El límite es que faltan los meses intermedios: se aprende variación estacional, pero no una secuencia diaria continua. Si las variables de congestión o rotación muestran potencial, conseguir más meses consecutivos será más valioso que repetir muchas veces el mismo muestreo.")

    add_page_break(doc)

    # Weather.
    add_small_label(doc, "3 · Meteorología")
    doc.add_heading("Por qué se ha aplazado, no descartado", level=1)
    add_body(doc, "El tiempo atmosférico puede explicar retrasos, pero añadirlo bien no consiste solo en unir dos tablas. Hay que asignar cada aeropuerto a una estación o rejilla, escoger la hora que realmente estaba disponible a T−60, controlar husos horarios, medir cobertura y decidir cómo tratar observaciones ausentes. Si esa unión se hace mal, se puede introducir información futura o crear una señal artificial.")

    doc.add_heading("Razones de la decisión", level=2)
    add_number(doc, "Establecer primero una referencia sólida con datos aeronáuticos. Sin baseline no se puede medir cuánto aporta realmente la meteorología.")
    add_number(doc, "Reducir complejidad mientras se valida el horizonte T−60, el split temporal y las reglas de fuga de información.")
    add_number(doc, "No gastar memoria y tiempo de cálculo en una fuente adicional antes de demostrar que el pipeline básico funciona.")
    add_number(doc, "Separar el efecto del modelo, de las variables operacionales y del tiempo meteorológico.")

    add_callout(doc, "Condición para incorporarla", "La meteorología debe entrar con datos históricos o pronósticos que hubieran estado disponibles una hora antes; nunca con la observación final conocida después del vuelo.")

    doc.add_heading("Qué se propondrá cuando se añada", level=2)
    add_bullet(doc, "Viento, visibilidad, precipitación, techo de nubes y fenómenos adversos en salida y llegada.")
    add_bullet(doc, "Unión por aeropuerto y hora, con tolerancias explícitas y auditoría de cobertura.")
    add_bullet(doc, "Ablación: comparar el mismo modelo con y sin meteorología, usando idénticos periodos y filas.")
    add_bullet(doc, "Medir la mejora en vuelos normales y en retrasos importantes, no solo el promedio global.")

    doc.add_heading("Conclusión sobre los datos", level=2)
    add_body(doc, "Antes de buscar otra fuente, el proyecto ya ha encontrado mejora con los vuelos anteriores observados a T−60. El siguiente escalado lógico es usar más meses consecutivos de vuelos; la meteorología será la siguiente familia externa de variables cuando el pipeline temporal esté congelado.")

    add_page_break(doc)

    # Initial analysis.
    add_small_label(doc, "4 · Análisis inicial")
    doc.add_heading("Qué dijeron las primeras exploraciones", level=1)
    add_body(doc, "Los análisis de las libretas 01 y 02 tuvieron tres funciones: comprobar que los datos podían leerse, localizar problemas de calidad y decidir qué variables merecían tratamiento especial. La muestra reproducible del 15% reunió 617.407 vuelos y 19 campos principales.")

    add_table(
        doc,
        ["Hallazgo", "Dato observado", "Decisión que provocó"],
        [
            ["Pocos nulos", "Registro 0,20%; coordenadas ≤0,06%; nivel solicitado 0,02% en la muestra.", "No eliminar columnas útiles; imputar o completar de forma controlada."],
            ["Cola larga de retrasos", "Mediana de llegada 2,47 min; percentil 75, 11,48 min; existen casos extremos.", "Usar mediana, MAE, P90 y segmentos; no depender solo de media/RMSE."],
            ["Valores físicamente dudosos", "Mínimos extremadamente negativos antes de limpiar.", "Fijar retraso mínimo de −120 min y auditar reglas."],
            ["Salida y llegada correlacionan", "Correlación aproximada 0,804.", "No usar el retraso de salida del mismo vuelo a T−60; sí reservarlo para post-off-block."],
            ["Alta variedad categórica", "1.169 ADEP, 1.181 ADES, 412 operadores y 230 tipos de avión en el alcance.", "One-hot solo para baja cardinalidad; hashing o categorías nativas para el resto."],
            ["Grupos operativos importan", "Ruta, operador, aeropuerto y aeronave muestran medias distintas.", "Crear baselines jerárquicos y variables históricas por entidad."],
        ],
        [2250, 3100, 4010],
        font_size=8.7,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 3. Del análisis exploratorio a decisiones concretas de modelado.")

    doc.add_heading("Cómo se interpretaron los outliers", level=2)
    add_body(doc, "La regla del rango intercuartílico marcaba aproximadamente un 4,42% de retrasos de llegada y un 5,55% de retrasos de salida como atípicos. No se eliminaron automáticamente: una demora grande puede ser un evento real y precisamente es un caso importante. El IQR se usó como diagnóstico; las eliminaciones se limitaron a reglas físicas y de alcance explicables.")

    doc.add_heading("Distribución de entrenamiento tras la limpieza", level=2)
    add_body(doc, "En train, la mediana del retraso de llegada es 2,42 minutos y la media 4,38. El 18,16% supera 15 minutos y el 1,05% supera 60. Esta asimetría explica por qué un modelo puede acertar bien la mayoría de vuelos y, al mismo tiempo, fallar en los retrasos relevantes.")
    add_callout(doc, "Implicación", "No basta con publicar un MAE global. Todas las comparaciones deben mostrar también el error en vuelos con más de 15 minutos y, como diagnóstico de cola, RMSE, mediana del error absoluto y percentil 90.")

    add_page_break(doc)

    # Cleaning.
    add_small_label(doc, "5 · Limpieza y preparación")
    doc.add_heading("Reglas aplicadas y razones", level=1)
    add_body(doc, "La limpieza se implementó con PySpark para procesar los 4,1 millones de filas de forma reproducible. De los datos brutos, 3.671.885 vuelos pertenecían al alcance regular programado y 3.671.258 quedaron después de las reglas físicas. Por tanto, casi toda la reducción (443.778 vuelos) procede de definir la población, no de borrar observaciones por comodidad.")

    add_table(
        doc,
        ["Regla", "Tratamiento", "Por qué"],
        [
            ["ICAO Flight Type", "Conservar solo S y retirar luego la columna.", "Define el alcance; después queda constante."],
            ["Retrasos", "Excluir valores inferiores a −120 min.", "Quita casos incoherentes sin cortar los grandes retrasos reales."],
            ["Requested FL", "Aceptar 0–500 o nulo; imputar mediana 340 aprendida en train.", "Mantiene valores plausibles y evita aprender de validación/test."],
            ["Distancia real", "Exigir >0 si existe, pero no usarla a T−60.", "Es una regla de calidad; la distancia real solo se conoce después."],
            ["Coordenadas", "Validar rangos y completar mediante dimensión de aeropuertos.", "La fuente auxiliar es más coherente que rellenar con constantes."],
            ["Categóricas nulas", "Usar Unknown cuando la variable se conserva.", "Evita perder filas y mantiene explícita la ausencia."],
            ["Tipo de aeronave raro", "Agrupar como OTHER si tiene <1.000 casos en train.", "Reduce ruido; el criterio se aprende solo en train."],
            ["Fechas", "Formato explícito; cero fallos de parseo en los cuatro tiempos.", "Evita interpretaciones regionales ambiguas."],
        ],
        [2050, 3400, 3910],
        font_size=8.45,
        first_col_bold=True,
    )

    doc.add_heading("Conteos finales y particiones", level=2)
    add_table(
        doc,
        ["Partición", "Periodo", "Filas limpias", "Uso"],
        [
            ["Train", "hasta sep-2022", "2.457.169", "Aprender imputaciones, categorías, agregados y modelos."],
            ["Validación", "dic-2022", "591.391", "Elegir decisiones e hiperparámetros."],
            ["Test", "mar-2023", "622.698", "Evaluación final después de congelar el modelo."],
        ],
        [1500, 1900, 1900, 4060],
        font_size=9,
    )

    add_callout(doc, "Transparencia sobre test", "La libreta de baselines abrió marzo de 2023 una vez después de congelar su ganador. Los modelos posteriores no han usado ese Parquet para ajustar decisiones. Para una medición final completamente ciega del modelo definitivo, conviene reservar un nuevo periodo posterior o declarar expresamente esta primera apertura.", color=ORANGE, fill="FFF5E6")

    add_page_break(doc)

    # Technology and categorical decisions.
    add_small_label(doc, "6 · Ingeniería de variables")
    doc.add_heading("PySpark, Parquet y categorías", level=1)

    doc.add_heading("Por qué PySpark", level=2)
    add_body(doc, "PySpark permite aplicar el mismo esquema, filtros, joins, agregaciones y transformaciones a millones de vuelos sin cargar todo en memoria como un único DataFrame de Pandas. Se ha usado en modo local en el ordenador, no en un clúster. No hay HDFS ni un despliegue Hadoop externo; solo el motor Spark y un pequeño adaptador de sistema de archivos para Windows.")

    doc.add_heading("Qué aporta Parquet", level=2)
    add_body(doc, "Parquet es un formato columnar. Guarda tipos de datos, comprime bien y permite leer solo las columnas necesarias. Aquí sirve para congelar train, validación y test después de la limpieza, evitando repetir cada vez la lectura y preparación de los CSV comprimidos. No es una base de datos ni implica usar Hadoop.")

    doc.add_heading("Decisiones por cardinalidad", level=2)
    add_table(
        doc,
        ["Variable", "Categorías en vuelos S", "Tratamiento", "Motivo"],
        [
            ["ICAO Flight Type", "1", "Eliminar", "Es siempre S tras filtrar."],
            ["STATFOR Market Segment", "6", "One-hot", "Muy pocas categorías; columnas claras e interpretables."],
            ["AC Type", "230", "OTHER + hashing; nativa en CatBoost", "One-hot sería ancho y sensible a tipos nuevos o raros."],
            ["ADEP", "1.169", "Hashing", "Alta cardinalidad."],
            ["ADES", "1.181", "Hashing", "Alta cardinalidad."],
            ["AC Operator", "412", "Hashing", "Cardinalidad moderada-alta y categorías nuevas."],
            ["Clase / motores del avión", "Baja", "One-hot", "Dimensiones agregadas y fáciles de interpretar."],
        ],
        [2150, 1700, 2600, 2910],
        font_size=8.55,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 4. One-hot se reserva para pocas categorías; AC Type no se codifica con one-hot.")

    doc.add_heading("Tamaño del hashing", level=2)
    add_body(doc, "Las cuatro variables enviadas a hashing suman 2.992 tokens distintos. El análisis teórico propone 32.768 posiciones: estima alrededor de un 4,4% de colisiones, frente a 8,6% con 16.384. En el benchmark Spark de memoria baja se redujo temporalmente a 8.192; en la Ridge T−60 se volvió a 32.768 usando una matriz dispersa. Antes del entrenamiento final debe congelarse una configuración tras comparar calidad, memoria y tiempo.")

    add_page_break(doc)

    # Baselines.
    add_small_label(doc, "7 · Baselines")
    doc.add_heading("Qué debe superar cualquier modelo", level=1)
    add_body(doc, "Un baseline es una regla sencilla que evita celebrar mejoras inexistentes. El primero predice para todos los vuelos la mediana del retraso del train: 2,42 minutos. Es robusto frente a valores extremos, pero ignora ruta, aeropuerto y aerolínea.")

    doc.add_heading("Baseline histórico de ruta + aerolínea", level=2)
    add_body(doc, "La versión más útil calcula medianas solo con train y usa una jerarquía: ruta + aerolínea → ruta → aeropuerto de salida + aerolínea → aeropuerto de salida → mediana global. Se compararon mínimos de 20, 50, 100 y 200 vuelos; 20 obtuvo el mejor MAE de validación. Para la ruta sin aerolínea se mantiene un mínimo de 100 vuelos.")

    add_table(
        doc,
        ["Baseline en validación", "MAE global", "MAE >15 min", "RMSE", "P90 abs."],
        [
            ["Mediana global", "11,685", "27,341", "17,885", "24,167"],
            ["Ruta con fallback", "10,135", "22,117", "15,923", "20,817"],
            ["Ruta + aerolínea con fallback", "9,953", "21,286", "15,669", "20,400"],
        ],
        [3000, 1420, 1600, 1420, 1920],
        font_size=9,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 5. Métricas sobre las 591.391 filas de validación de la libreta 05.")

    doc.add_heading("Cobertura del fallback ganador", level=2)
    add_table(
        doc,
        ["Nivel usado", "Vuelos de validación", "Porcentaje"],
        [
            ["Ruta + aerolínea", "552.210", "93,37%"],
            ["Aeropuerto salida + aerolínea", "22.918", "3,88%"],
            ["Ruta", "10.171", "1,72%"],
            ["Aeropuerto de salida", "5.135", "0,87%"],
            ["Mediana global", "957", "0,16%"],
        ],
        [3900, 2730, 2730],
        font_size=9.2,
    )
    add_body(doc, "En test, ya con el baseline congelado, la jerarquía obtuvo MAE 10,18; el error subió frente a validación, algo esperable cuando cambia el periodo. Este baseline sigue siendo una referencia muy competitiva y fácil de explicar.")

    add_callout(doc, "Baseline post-off-block", "Usar el retraso de salida para actualizar la llegada sería operacionalmente lógico después del off-block, pero no se ha mezclado con la tarea T−60. Tendrá su propio benchmark cuando se retome ese horizonte.")

    add_page_break(doc)

    # First models.
    add_small_label(doc, "8 · Primeros modelos")
    doc.add_heading("Qué se probó y qué se aprendió", level=1)
    add_body(doc, "Para poder experimentar con poca RAM, la libreta 06 usó una muestra determinista del 1% de train (24.743 filas) y el 5% de validación (29.315). Esto sirve para comprobar código, memoria y dirección de las mejoras, pero no equivale al entrenamiento final con todos los datos.")

    add_table(
        doc,
        ["Modelo / variante", "MAE global", "MAE >15 min", "Puntuación combinada", "Lectura"],
        [
            ["Lineal regularizada · original", "10,871", "19,268", "15,069", "Mejor equilibrio inicial."],
            ["Lineal · log", "10,876", "19,323", "15,099", "Sin mejora."],
            ["Lineal · Yeo–Johnson", "10,874", "19,338", "15,106", "Sin mejora."],
            ["XGBoost", "10,676", "23,199", "16,938", "Mejor global que lineal; peor en retrasados."],
            ["Gradient-Boosted Trees", "10,924", "24,043", "17,484", "No supera el equilibrio lineal."],
            ["Random Forest", "11,569", "24,577", "18,073", "Peor en este presupuesto."],
        ],
        [2700, 1300, 1450, 1700, 2210],
        font_size=8.55,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 6. Resultados de validación alineada; 0,5 × MAE global + 0,5 × MAE de vuelos >15 min.")

    doc.add_heading("Decisión sobre log y Yeo–Johnson", level=2)
    add_body(doc, "Ambas transformaciones siguen disponibles como parámetros opcionales de las funciones. Se ajustan únicamente con train para no filtrar información. En la primera comparación no mejoraron la versión original, por lo que no se activan por defecto. Esto no significa que sean incorrectas: significa que no aportaron valor medible en este conjunto y configuración.")

    doc.add_heading("Por qué XGBoost fue el cuarto modelo", level=2)
    add_body(doc, "XGBoost ofrece árboles potenciados, interacciones y no linealidad con una integración viable en el entorno actual. LightGBM mediante SynapseML quedó como experimento futuro porque las versiones publicadas de SynapseML estaban orientadas a Spark 3.x / Scala 2.12, mientras el proyecto usa Spark 4.2 / Scala 2.13. Introducirlo ahora habría debilitado la reproducibilidad.")

    add_callout(doc, "Lectura correcta", "Los árboles no están descartados. Sus resultados dicen que, con pocas filas y variables todavía básicas, aumentar complejidad no bastó. La siguiente prueba razonable fue un algoritmo especialmente fuerte con categorías: CatBoost.")

    add_page_break(doc)

    # CatBoost and aligned results.
    add_small_label(doc, "9 · CatBoost")
    doc.add_heading("Por qué probarlo y qué ocurrió", level=1)
    add_body(doc, "CatBoost está diseñado para trabajar bien con variables categóricas y aprender interacciones no lineales sin convertir cada categoría en cientos de columnas one-hot. Esto encaja con aeropuertos, operador y tipo de aeronave. También permite respetar el orden temporal durante el entrenamiento mediante has_time=True.")

    doc.add_heading("Curva de aprendizaje", level=2)
    add_table(
        doc,
        ["Train usado", "Filas", "MAE global", "MAE >15 min", "Combinado"],
        [
            ["1%", "24.743", "10,370", "21,030", "15,700"],
            ["5%", "123.135", "10,059", "20,601", "15,330"],
            ["10%", "245.590", "9,984", "20,484", "15,234"],
        ],
        [1750, 1750, 1900, 1900, 2060],
        font_size=9.2,
    )
    add_body(doc, "Pasar del 1% al 5% mejora la puntuación combinada en 0,37 minutos; del 5% al 10%, solo 0,10. Hay mejora con más datos, pero con rendimientos decrecientes. Esto justificó seguir escalando de forma gradual y, a la vez, invertir en variables más informativas.")

    doc.add_heading("Por qué no se ponderaron agresivamente los retrasados", level=2)
    add_body(doc, "Se ensayaron pesos para dar más importancia a vuelos con más de 15 minutos. El ajuste agresivo redujo su MAE a 13,68, pero elevó el MAE global a 16,23. El modelo casi dejó de representar al conjunto completo. Se rechazó mediante una regla previa: no aceptar una mejora de un segmento si degrada demasiado el resultado global.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(model_chart), width=Inches(6.5))
    add_caption(doc, "Figura 2. Comparación orientativa. Los presupuestos de entrenamiento no son idénticos; las cifras deben leerse junto a su alcance.")

    add_callout(doc, "Conclusión de CatBoost", "Es un competidor fuerte y mejora el error global, pero antes de las nuevas variables operacionales no superaba a la lineal regularizada en la puntuación equilibrada.")

    add_page_break(doc)

    # Operational features.
    add_small_label(doc, "10 · Variables operacionales T−60")
    doc.add_heading("Usar solo vuelos anteriores ya observados", level=1)
    add_body(doc, "La libreta 08 construye señales de estado de la red justo antes de cada predicción. Para cada vuelo objetivo, el corte es su off-block programado menos 60 minutos. Una salida histórica solo puede contar si su off-block real ya ocurrió; una llegada histórica, si su llegada real ya ocurrió. Además se retiran contribuciones del mismo vuelo cuando una salida excepcionalmente temprana podría contarse a sí misma.")

    doc.add_heading("Variables creadas", level=2)
    add_table(
        doc,
        ["Entidad", "Ventanas", "Medidas", "Interpretación"],
        [
            ["Aeropuerto de salida", "1, 6 y 24 h", "conteo, media, desviación, proporción >15", "Congestión y retraso reciente en salidas."],
            ["Aeropuerto de llegada", "1, 6 y 24 h", "conteo, media, desviación, proporción >15", "Estado reciente del destino."],
            ["Ruta", "1, 6 y 24 h", "conteo, media, desviación, proporción >15", "Comportamiento reciente del trayecto."],
            ["Operador", "1, 6 y 24 h", "salidas y llegadas observadas", "Propagación operacional de la aerolínea."],
            ["Rotación del avión", "último vuelo completado", "retraso previo y antigüedad", "Efecto de la aeronave que continúa su secuencia."],
        ],
        [2100, 1500, 2650, 3110],
        font_size=8.65,
        first_col_bold=True,
    )

    doc.add_heading("Auditoría de fuga", level=2)
    add_bullet(doc, "Cero eventos históricos posteriores al corte T−60.")
    add_bullet(doc, "Se eliminaron 216 autocontaminaciones en ADEP y las mismas 216 en operador, debidas a vuelos que salieron más de una hora antes de lo programado.")
    add_bullet(doc, "Marzo de 2023 no se incorporó al histórico de construcción de variables.")
    add_bullet(doc, "La matrícula estaba informada en el 99,987% de la muestra; se asume disponible a T−60 y esta condición debe verificarse en producción.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(ablation_chart), width=Inches(5.85))
    add_caption(doc, "Figura 3. Ablación interna con Ridge: se conservan todas las entidades y la rotación, pero se elimina la ventana de 6 horas.")

    add_body(doc, "La selección final conserva 51 variables numéricas, las ventanas de 1 y 24 horas y la rotación. La ventana de 6 horas duplicaba información intermedia sin mejorar la puntuación combinada.")

    add_page_break(doc)

    # Final current models.
    add_small_label(doc, "11 · Resultado actual")
    doc.add_heading("Ridge frente a CatBoost", level=1)
    add_body(doc, "Las variables y los hiperparámetros se eligieron en un tuning temporal interno: 172.315 filas para ajuste y 73.275 para tuning. Después se reentrenó con el 10% completo de train y se abrió una única validación de 29.315 filas. Test no se leyó en esta etapa.")

    add_table(
        doc,
        ["Candidato T−60", "MAE global", "MAE >15 min", "Combinado", "Conclusión"],
        [
            ["Ridge + variables operacionales", "9,941", "18,207", "14,074", "Mejor equilibrio; candidato actual."],
            ["Ensemble 70% Ridge + 30% CatBoost", "9,716", "18,528", "14,122", "Mejor global que Ridge; no gana en combinado."],
            ["CatBoost + variables operacionales", "9,537", "19,635", "14,586", "Mejor MAE global; peor en retrasados."],
            ["Mejor candidato previo: lineal original", "10,871", "19,268", "15,069", "Referencia antes de variables operacionales."],
        ],
        [2750, 1350, 1450, 1450, 2360],
        font_size=8.55,
        first_col_bold=True,
    )
    add_caption(doc, "Tabla 7. Comparación final actual. La puntuación combinada pesa 50% MAE global y 50% MAE en retrasos >15 min.")

    add_callout(doc, "Mejora observada", "Ridge reduce la puntuación combinada en 0,995 minutos frente al mejor candidato anterior; el MAE global mejora en 0,931 y el de retrasados en 1,060 minutos. Se supera el criterio fijado para escalar más allá del 10%.")

    doc.add_heading("Por qué Ridge", level=2)
    add_bullet(doc, "Regulariza: reduce coeficientes extremos y funciona bien con muchas variables correlacionadas y vectores dispersos de hashing.")
    add_bullet(doc, "Es estable y barato en memoria, importante en un ordenador con alrededor de 5–6 GB libres durante las ejecuciones.")
    add_bullet(doc, "Es explicable: permite estudiar el signo y tamaño de los efectos tras controlar la codificación.")
    add_bullet(doc, "Ha obtenido el mejor compromiso acordado entre el vuelo típico y los retrasos importantes.")

    doc.add_heading("Por qué mantener CatBoost", level=2)
    add_bullet(doc, "Trata categorías de forma nativa y aprende interacciones no lineales entre ruta, operador, avión, horario y estado reciente.")
    add_bullet(doc, "Obtiene el mejor MAE global actual: 9,54 minutos.")
    add_bullet(doc, "Puede aportar diversidad a un ensemble o modelar residuos de Ridge/baseline.")
    add_bullet(doc, "Es más costoso y, en la configuración actual, pierde precisión en vuelos retrasados; por eso no es el ganador operativo todavía.")

    add_page_break(doc)

    # Challenges.
    add_small_label(doc, "12 · Retos y soluciones")
    doc.add_heading("Dificultades encontradas y cómo se resolvieron", level=1)
    add_table(
        doc,
        ["Reto", "Solución aplicada", "Por qué fue apropiada"],
        [
            ["Memoria limitada", "Parquet, selección de columnas, hashing disperso, muestras deterministas 1/5/10% y modelos secuenciales.", "Permite experimentar sin bloquear el equipo y conservar reproducibilidad."],
            ["Entorno Spark en Windows", "Sesión local controlada, pocas particiones y sin depender de un clúster Hadoop.", "Reduce dependencias y hace el proyecto ejecutable en el portátil."],
            ["Narrativa y notebooks cambiantes", "Restaurar la libreta categórica y centralizar reglas en flight_config.py.", "Conserva análisis previo sin perder las decisiones nuevas."],
            ["Fuga temporal", "Lista de variables por horizonte, split temporal, fit solo en train y auditoría evento a evento a T−60.", "Convierte el rendimiento medido en una expectativa operacional creíble."],
            ["Categorías nuevas y raras", "OTHER, hashing y fallback jerárquico; CatBoost nativo como alternativa.", "Evita columnas enormes y predicciones rotas por valores no vistos."],
            ["Métrica engañosa", "MAE global + MAE >15 min, RMSE, MedAE y P90; guardrail global.", "Impide mejorar un segmento destruyendo el resto."],
            ["Comparaciones con presupuestos distintos", "Declarar porcentajes y filas; usar validación alineada cuando es posible.", "Evita afirmar que un algoritmo gana por una comparación injusta."],
        ],
        [2300, 3850, 3210],
        font_size=8.45,
        first_col_bold=True,
    )

    doc.add_heading("Qué no se debe concluir todavía", level=2)
    add_bullet(doc, "Ridge no es el modelo definitivo: su ventaja se ha medido con 10% de train y una muestra de validación.")
    add_bullet(doc, "CatBoost no ha perdido de forma general: es el mejor en MAE global y puede mejorar con ajuste o más datos.")
    add_bullet(doc, "El error en retrasados no implica que sepamos de antemano qué vuelos se retrasarán; el segmento solo se usa después para evaluar.")
    add_bullet(doc, "Las ventanas de vuelos previos son válidas solo si la infraestructura de producción puede reconstruirlas con datos disponibles en tiempo real.")

    add_page_break(doc)

    # Next steps.
    add_small_label(doc, "13 · Mejoras propuestas")
    doc.add_heading("Plan recomendado a partir de aquí", level=1)
    add_body(doc, "El siguiente objetivo no debería ser probar modelos al azar, sino confirmar que la mejora se mantiene al crecer el volumen y al cambiar el periodo. La secuencia recomendada es la siguiente.")

    add_table(
        doc,
        ["Prioridad", "Acción", "Criterio para continuar"],
        [
            ["1", "Entrenar Ridge y CatBoost con 25%; después 50% y 100% solo si la curva sigue mejorando.", "Ganancia estable en combinado y sin degradación global relevante."],
            ["2", "Usar varios cortes temporales de validación, no un único mes.", "Resultados consistentes entre estaciones y niveles de tráfico."],
            ["3", "Congelar pipeline, variables, hash y parámetros; evaluar en un holdout nuevo posterior.", "Estimación final sin decisiones posteriores sobre ese periodo."],
            ["4", "Probar clasificación (>15 min) + regresión de minutos, o CatBoost sobre residuos.", "Mejorar retrasados sin romper el MAE global."],
            ["5", "Conseguir más meses consecutivos de vuelos.", "Mejor cobertura para congestión, rotaciones y cambios operativos."],
            ["6", "Añadir meteorología con unión estrictamente T−60 y ablación.", "Mejora atribuible y cero observaciones futuras."],
            ["7", "Calibrar y explicar errores por ruta, operador, aeropuerto, estación y severidad.", "Detectar sesgos y casos donde el modelo no es confiable."],
        ],
        [1150, 5280, 2930],
        font_size=8.55,
        first_col_bold=True,
    )

    doc.add_heading("Ajustes concretos de modelos", level=2)
    add_bullet(doc, "Ridge: comparar 16.384 y 32.768 posiciones de hashing, revisar estabilidad de coeficientes y mantener alpha=10 como punto de partida.")
    add_bullet(doc, "CatBoost: probar 1.200–1.500 iteraciones, learning rate 0,02–0,03 y early stopping en cortes temporales internos.")
    add_bullet(doc, "Ensemble: reoptimizar pesos en varios periodos; el 70/30 actual no generalizó mejor que Ridge en la puntuación combinada.")
    add_bullet(doc, "Baseline como variable: usar codificaciones históricas calculadas sin fuga y modelar el residuo respecto a ruta + aerolínea.")
    add_bullet(doc, "LightGBM/SynapseML: mantener como TODO en un entorno Spark compatible y aislado.")

    doc.add_heading("¿CNN-LSTM ahora?", level=2)
    add_body(doc, "No es la siguiente prioridad. CNN/LSTM cobra más sentido cuando hay secuencias densas y continuas por aeropuerto, ruta o aeronave. Los datos actuales son principalmente tabulares y seis cortes mensuales. Primero conviene conseguir meses consecutivos, construir secuencias sin huecos y comparar contra Ridge/CatBoost con el mismo split. Solo entonces se justificaría su mayor coste y complejidad.")

    add_callout(doc, "Decisión recomendada", "Entrenar primero con más datos de vuelos ya disponibles y validar en varios periodos. Añadir después meteorología. Probar una red secuencial solo cuando exista un histórico continuo suficiente y un baseline tabular congelado.")

    add_page_break(doc)

    # Conclusions and glossary.
    add_small_label(doc, "14 · Conclusión")
    doc.add_heading("Situación del proyecto", level=1)
    add_body(doc, "El proyecto ya dispone de una definición operacional clara, un pipeline reproducible, reglas de calidad justificadas, particiones temporales, baselines fuertes y una comparación inicial de modelos. La decisión de ampliar los vuelos antes de añadir meteorología permitió encontrar una señal útil en la propia operación: el estado reciente de aeropuertos, rutas, operadores y la rotación del avión.")
    add_body(doc, "Ridge es hoy la elección más razonable si se necesita un modelo sencillo, estable y equilibrado. CatBoost debe mantenerse porque obtiene el menor error global y puede capturar relaciones que Ridge no ve. La conclusión no es escoger uno para siempre, sino escalar ambos bajo el mismo protocolo y exigir que cualquier mejora se mantenga en periodos futuros.")
    add_callout(doc, "Mensaje final", "Con la evidencia actual, la mejor inversión es calidad temporal de las variables y más cobertura histórica. La complejidad del algoritmo debe añadirse solo cuando produzca una mejora reproducible y operativamente válida.")

    doc.add_heading("Glosario breve", level=2)
    add_table(
        doc,
        ["Término", "Explicación sencilla"],
        [
            ["T−60", "Momento situado 60 minutos antes de la salida programada."],
            ["Leakage", "Uso accidental de información que todavía no existiría al hacer la predicción."],
            ["MAE", "Error absoluto medio, en minutos. Menor es mejor."],
            ["RMSE", "Métrica que penaliza con fuerza los errores muy grandes."],
            ["MedAE", "Mediana del error absoluto; representa un vuelo típico."],
            ["P90", "El 90% de los errores queda por debajo de ese valor."],
            ["One-hot", "Una columna binaria por categoría; útil cuando hay pocas."],
            ["Hashing", "Proyección de muchas categorías a un vector de tamaño fijo; puede haber colisiones."],
            ["Ridge", "Regresión lineal regularizada que limita coeficientes extremos."],
            ["CatBoost", "Árboles potenciados preparados para variables categóricas."],
            ["Fallback", "Regla de respaldo cuando no hay suficiente historial para un grupo concreto."],
        ],
        [2200, 7160],
        font_size=8.75,
        first_col_bold=True,
    )

    doc.add_heading("Fuentes internas revisadas", level=2)
    add_body(doc, "Libretas 01_initial_analysis_sample a 08_arrival_pre_t60_operational_features; configuración compartida en src/flight_config.py; pipeline Spark en src/spark_flight_pipeline.py; salidas ejecutadas y gráficos de la carpeta reports. Las cifras de este informe proceden de esas ejecuciones locales.")

    # Prevent accidental orphan lines in all normal paragraphs where practical.
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {"Normal", "List Bullet", "List Bullet 2", "List Number"}:
            paragraph.paragraph_format.widow_control = True

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(path)
